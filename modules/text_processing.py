# modules/text_processing.py

from pathlib import Path
import json
from docx import Document
from typing import Any, List

class DocumentConverter:
    """
    Converts JSON-extracted data to DOCX or TXT documents.
    """
    def __init__(self, schema_name: str) -> None:
        self.schema_name: str = schema_name.lower()

    def extract_entries(self, json_file: Path) -> List[Any]:
        try:
            with json_file.open("r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            print(f"Error reading JSON file {json_file}: {e}")
            return []
        entries: List[Any] = []
        if isinstance(data, dict):
            if "entries" in data:
                entries = data["entries"]
            elif "responses" in data:
                for resp in data["responses"]:
                    try:
                        body = resp.get("body", {})
                        choices = body.get("choices", [])
                        if choices:
                            message = choices[0].get("message", {})
                            content = message.get("content", "")
                            content_json = json.loads(content)
                            if isinstance(content_json, dict) and "entries" in content_json:
                                entries.extend(content_json["entries"])
                    except Exception as e:
                        print(f"Error processing response: {e}")
        elif isinstance(data, list):
            for item in data:
                response = item.get("response")
                if isinstance(response, str):
                    try:
                        response_obj = json.loads(response)
                        if isinstance(response_obj, dict) and "entries" in response_obj:
                            entries.extend(response_obj["entries"])
                    except Exception as e:
                        print(f"Error parsing response string: {e}")
                elif isinstance(response, dict) and "entries" in response:
                    entries.extend(response["entries"])
        return entries

    def convert_to_docx(self, json_file: Path, output_file: Path) -> None:
        entries: List[Any] = self.extract_entries(json_file)
        document: Document = Document()
        document.add_heading(json_file.stem, 0)
        converters = {
            "structuredsummaries": self._convert_structured_summaries_to_docx,
            "bibliographicentries": self._convert_bibliographic_entries_to_docx,
            "historicaladdressbookentries": self._convert_historicaladdressbookentries_to_docx,
            "brazilianoccupationrecords": self._convert_brazilianoccupationrecords_to_docx
        }
        key = self.schema_name.lower()
        if key in converters:
            converters[key](entries, document)
        else:
            for entry in entries:
                document.add_paragraph(str(entry))
        try:
            document.save(output_file)
            print(f"DOCX file generated at {output_file}")
        except Exception as e:
            print(f"Error saving DOCX file {output_file}: {e}")

    def convert_to_txt(self, json_file: Path, output_file: Path) -> None:
        entries: List[Any] = self.extract_entries(json_file)
        lines: List[str] = []
        converters = {
            "structuredsummaries": self._convert_structured_summaries_to_txt,
            "bibliographicentries": self._convert_bibliographic_entries_to_txt,
            "historicaladdressbookentries": self._convert_historicaladdressbookentries_to_txt,
            "brazilianoccupationrecords": self._convert_brazilianoccupationrecords_to_txt
        }
        key = self.schema_name.lower()
        if key in converters:
            lines = converters[key](entries)
        else:
            for entry in entries:
                lines.append(str(entry))
        try:
            with output_file.open("w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            print(f"TXT file generated at {output_file}")
        except Exception as e:
            print(f"Error saving TXT file {output_file}: {e}")

    # --- Schema-Specific DOCX Converters ---
    def _convert_structured_summaries_to_docx(self, entries: list,
                                              document: "Document") -> None:
        """
        Converts structured summaries entries to a DOCX document.
        For each entry, writes the page number in bold (as simple text) followed by bullet-pointed summaries.
        The keywords are formatted in italic (without asterisks).
        """
        literature_set = set()
        for entry in entries:
            page = entry.get("page")
            bullet_points = entry.get("bullet_points")
            keywords = entry.get("keywords")
            literature = entry.get("literature")

            # Instead of adding a heading, add a paragraph with bold text for the page number.
            p_page = document.add_paragraph()
            run_page = p_page.add_run(
                f"Page {page}" if page is not None else "Page Unknown")
            run_page.bold = True

            # For keywords: add a bullet point with italic text, without asterisks.
            if keywords and isinstance(keywords, list):
                formatted_keywords = ", ".join(kw for kw in keywords if kw)
                p_keyword = document.add_paragraph(style='List Bullet')
                p_keyword.add_run("Keywords: ")
                run_keywords = p_keyword.add_run(formatted_keywords)
                run_keywords.italic = True

            # Add each bullet point as a separate bullet item.
            if bullet_points and isinstance(bullet_points, list):
                for bp in bullet_points:
                    if bp:
                        p_bp = document.add_paragraph(style='List Bullet')
                        p_bp.add_run(str(bp))
            else:
                document.add_paragraph("No bullet points available.")

            # Add literature references if available.
            if literature and isinstance(literature, list):
                formatted_refs = ", ".join(
                    str(ref) for ref in literature if ref)
                p_ref = document.add_paragraph(style='List Bullet')
                p_ref.add_run(f"References: {formatted_refs}")

            # Accumulate literature for the consolidated literature section.
            if literature and isinstance(literature, list):
                for lit in literature:
                    if lit:
                        literature_set.add(lit)

            # Add an empty paragraph for spacing.
            document.add_paragraph("")

        # If any literature was found, add a final section.
        if literature_set:
            document.add_page_break()
            document.add_heading("Literature", level=1)
            for lit in sorted(literature_set):
                document.add_paragraph(str(lit), style='List Bullet')

    def _convert_bibliographic_entries_to_docx(self, entries: List[Any],
                                               document: Document) -> None:
        for entry in entries:
            # Extract primary entry data
            full_title = entry.get("full_title", "Unknown Title")
            short_title = entry.get("short_title", "")
            culinary_focus = entry.get("culinary_focus", [])
            book_format = entry.get("format", "")
            pages = entry.get("pages", "")
            edition_info = entry.get("edition_info", [])
            total_editions = entry.get("total_editions", "")

            # Add entry header and basic information
            document.add_heading(full_title, level=1)
            document.add_paragraph(f"Short Title: {short_title}")

            # Add culinary focus areas
            if culinary_focus and isinstance(culinary_focus, list):
                document.add_paragraph(
                    f"Culinary Focus: {', '.join(culinary_focus)}")
            else:
                document.add_paragraph("Culinary Focus: Unknown")

            # Add format and page information
            document.add_paragraph(
                f"Format: {book_format if book_format else 'Unknown'}")
            document.add_paragraph(
                f"Pages: {pages if pages is not None else 'Unknown'}")
            document.add_paragraph(
                f"Total Editions: {total_editions if total_editions is not None else 'Unknown'}")

            # Add edition information
            document.add_heading("Edition Information", level=2)
            if edition_info and isinstance(edition_info, list):
                for edition in edition_info:
                    # Extract edition data
                    edition_number = edition.get("edition_number", "Unknown")
                    year = edition.get("year", "Unknown")

                    # Extract location data
                    location = edition.get("location", {})
                    country = location.get("country",
                                           "Unknown") if location else "Unknown"
                    city = location.get("city",
                                        "Unknown") if location else "Unknown"

                    # Get contributors
                    contributors = edition.get("contributors", [])
                    contributors_str = ", ".join(
                        contributors) if contributors and isinstance(
                        contributors, list) else "Unknown"

                    # Get other edition details
                    language = edition.get("language", "")
                    short_note = edition.get("short_note", "")
                    edition_category = edition.get("edition_category", "")

                    # Format edition text
                    edition_text = (
                        f"Edition: {edition_number if edition_number is not None else 'Unknown'}, "
                        f"Year: {year if year is not None else 'Unknown'}, "
                        f"Location: {city}, {country}, "
                        f"Language: {language if language else 'Unknown'}, "
                        f"Contributors: {contributors_str}, "
                        f"Category: {edition_category if edition_category else 'Unknown'}"
                    )

                    if short_note:
                        edition_text += f"\nNote: {short_note}"

                    document.add_paragraph(edition_text, style='List Bullet')
            else:
                document.add_paragraph("No edition information available.",
                                       style='List Bullet')

            document.add_page_break()

    def _convert_historicaladdressbookentries_to_docx(self, entries: List[Any],
                                                      document: Document) -> None:
        for entry in entries:
            last_name = entry.get("last_name", "Unknown")
            first_name = entry.get("first_name", "Unknown")
            occupation = entry.get("occupation", "Unknown")
            section = entry.get("section")
            honorific = entry.get("honorific")
            additional_notes = entry.get("additional_notes")

            header = f"{last_name}, {first_name} - {occupation}"
            if section:
                header += f" (Section: {section})"
            document.add_heading(header, level=1)

            # Extract and add address information.
            address = entry.get("address", {})
            street = address.get("street", "Unknown")
            street_number = address.get("street_number", "*0*")
            document.add_paragraph(f"Address: {street} {street_number}")

            details = []
            if honorific:
                details.append(f"Honorific: {honorific}")
            if additional_notes:
                details.append(f"Additional Notes: {additional_notes}")
            if details:
                document.add_paragraph("; ".join(details))
            document.add_page_break()

    def _convert_brazilianoccupationrecords_to_docx(self, entries: List[Any], document: Document) -> None:
        for entry in entries:
            header = f"{entry.get('surname', '')}, {entry.get('first_name', '')} - {entry.get('profession', '')}"
            document.add_heading(header, level=1)
            document.add_paragraph(f"Record Header: {entry.get('record_header', '')}")
            document.add_paragraph(f"Location: {entry.get('location', '')}")
            document.add_paragraph(f"Height: {entry.get('height', '')}")
            document.add_paragraph(f"Skin Color: {entry.get('skin_color', '')}")
            document.add_paragraph(f"Hair Color: {entry.get('hair_color', '')}")
            document.add_paragraph(f"Hair Texture: {entry.get('hair_texture', '')}")
            document.add_paragraph(f"Beard: {entry.get('beard', '')}")
            document.add_paragraph(f"Mustache: {entry.get('mustache', '')}")
            document.add_paragraph(f"Assignatura: {entry.get('assignatura', '')}")
            document.add_paragraph(f"Reservista: {entry.get('reservista', '')}")
            document.add_paragraph(f"Eyes: {entry.get('eyes', '')}")
            document.add_paragraph(f"Mouth: {entry.get('mouth', '')}")
            document.add_paragraph(f"Face: {entry.get('face', '')}")
            document.add_paragraph(f"Nose: {entry.get('nose', '')}")
            document.add_paragraph(f"Marks: {entry.get('marks', '')}")
            officials = entry.get("officials", [])
            if officials:
                officials_str = "; ".join(
                    [f"{o.get('position', '')}: {o.get('signature', '')}" for o in officials]
                )
            else:
                officials_str = ""
            document.add_paragraph(f"Officials: {officials_str}")
            document.add_paragraph(f"Father: {entry.get('father', '')}")
            document.add_paragraph(f"Mother: {entry.get('mother', '')}")
            document.add_paragraph(f"Birth Date: {entry.get('birth_date', '')}")
            document.add_paragraph(f"Birth Place: {entry.get('birth_place', '')}")
            document.add_paragraph(f"Municipality: {entry.get('municipality', '')}")
            document.add_paragraph(f"Civil Status: {entry.get('civil_status', '')}")
            document.add_paragraph(f"Vaccinated: {entry.get('vaccinated', '')}")
            document.add_paragraph(f"Can Read: {entry.get('can_read', '')}")
            document.add_paragraph(f"Can Write: {entry.get('can_write', '')}")
            document.add_paragraph(f"Can Count: {entry.get('can_count', '')}")
            document.add_paragraph(f"Swimming: {entry.get('swimming', '')}")
            document.add_paragraph(f"Cyclist: {entry.get('cyclist', '')}")
            document.add_paragraph(f"Motorcyclist: {entry.get('motorcyclist', '')}")
            document.add_paragraph(f"Driver: {entry.get('driver', '')}")
            document.add_paragraph(f"Chauffeur: {entry.get('chauffeur', '')}")
            document.add_paragraph(f"Telegraphist: {entry.get('telegraphist', '')}")
            document.add_paragraph(f"Telephonist: {entry.get('telephonist', '')}")
            document.add_paragraph(f"Residence: {entry.get('residence', '')}")
            document.add_paragraph(f"Observations: {entry.get('observations', '')}")
            document.add_page_break()

    # --- Schema-Specific TXT Converters ---
    def _convert_structured_summaries_to_txt(self,
                                             entries: List[Any]) -> List[str]:
        lines: List[str] = []
        literature_set = set()
        for entry in entries:
            page = entry.get("page", "Unknown")
            bullet_points = entry.get("bullet_points", [])
            keywords = entry.get("keywords", [])
            literature = entry.get("literature", [])
            lines.append(f"Page {page}")
            # Add keywords bullet point.
            if keywords and isinstance(keywords, list):
                formatted_keywords = ", ".join(f"*{kw}*" for kw in keywords if kw)
                lines.append(f" - Keywords: {formatted_keywords}")
            # List the bullet points.
            if bullet_points and isinstance(bullet_points, list):
                for bp in bullet_points:
                    if bp:
                        lines.append(f" - {bp}")
            else:
                lines.append("No bullet points available.")
            # Append a bullet point for literature references.
            if literature and isinstance(literature, list):
                formatted_refs = ", ".join(str(ref) for ref in literature if ref)
                lines.append(f" - References: {formatted_refs}")
            # Collect literature for the final consolidated section.
            if literature and isinstance(literature, list):
                for lit in literature:
                    if lit:
                        literature_set.add(lit)
            lines.append("")
        if literature_set:
            lines.append("Literature:")
            for lit in sorted(literature_set):
                lines.append(f" - {lit}")
        return lines

    def _convert_bibliographic_entries_to_txt(self, entries: List[Any]) -> List[
        str]:
        lines: List[str] = []

        for entry in entries:
            # Extract primary entry data
            full_title = entry.get("full_title", "Unknown Title")
            short_title = entry.get("short_title", "")
            culinary_focus = entry.get("culinary_focus", [])
            book_format = entry.get("format", "")
            pages = entry.get("pages", "")
            edition_info = entry.get("edition_info", [])
            total_editions = entry.get("total_editions", "")

            # Add entry header and basic information
            lines.append(f"TITLE: {full_title}")
            lines.append(f"Short Title: {short_title}")

            # Add culinary focus areas
            if culinary_focus and isinstance(culinary_focus, list):
                lines.append(f"Culinary Focus: {', '.join(culinary_focus)}")
            else:
                lines.append("Culinary Focus: Unknown")

            # Add format and page information
            lines.append(f"Format: {book_format if book_format else 'Unknown'}")
            lines.append(f"Pages: {pages if pages is not None else 'Unknown'}")
            lines.append(
                f"Total Editions: {total_editions if total_editions is not None else 'Unknown'}")

            # Add edition information
            lines.append("\nEDITION INFORMATION:")
            if edition_info and isinstance(edition_info, list):
                for i, edition in enumerate(edition_info, 1):
                    # Extract edition data
                    edition_number = edition.get("edition_number", "Unknown")
                    year = edition.get("year", "Unknown")

                    # Extract location data
                    location = edition.get("location", {})
                    country = location.get("country",
                                           "Unknown") if location else "Unknown"
                    city = location.get("city",
                                        "Unknown") if location else "Unknown"

                    # Get contributors
                    contributors = edition.get("contributors", [])
                    contributors_str = ", ".join(
                        contributors) if contributors and isinstance(
                        contributors, list) else "Unknown"

                    # Get other edition details
                    language = edition.get("language", "")
                    short_note = edition.get("short_note", "")
                    edition_category = edition.get("edition_category", "")

                    # Format edition text
                    lines.append(f"  Edition {i}:")
                    lines.append(
                        f"    Edition Number: {edition_number if edition_number is not None else 'Unknown'}")
                    lines.append(
                        f"    Year: {year if year is not None else 'Unknown'}")
                    lines.append(f"    Location: {city}, {country}")
                    lines.append(
                        f"    Language: {language if language else 'Unknown'}")
                    lines.append(f"    Contributors: {contributors_str}")
                    lines.append(
                        f"    Category: {edition_category if edition_category else 'Unknown'}")

                    if short_note:
                        lines.append(f"    Note: {short_note}")

                    lines.append("")
            else:
                lines.append("  No edition information available.")

            lines.append("=" * 50)
            lines.append("")

        return lines

    def _convert_historicaladdressbookentries_to_txt(self,
                                                     entries: List[Any]) -> \
        List[str]:
        lines: List[str] = []
        for entry in entries:
            last_name = entry.get("last_name", "Unknown")
            first_name = entry.get("first_name", "Unknown")
            occupation = entry.get("occupation", "Unknown")
            section = entry.get("section")
            honorific = entry.get("honorific")
            additional_notes = entry.get("additional_notes")

            header = f"{last_name}, {first_name} - {occupation}"
            if section:
                header += f" (Section: {section})"
            lines.append(header)

            # Extract and append address information.
            address = entry.get("address", {})
            street = address.get("street", "Unknown")
            street_number = address.get("street_number", "*0*")
            lines.append(f"Address: {street} {street_number}")

            if honorific:
                lines.append(f"Honorific: {honorific}")
            if additional_notes:
                lines.append(f"Additional Notes: {additional_notes}")
            lines.append("")
        return lines

    def _convert_brazilianoccupationrecords_to_txt(self, entries: List[Any]) -> List[str]:
        lines: List[str] = []
        for entry in entries:
            header = f"{entry.get('surname', '')}, {entry.get('first_name', '')} - {entry.get('profession', '')}"
            lines.append(header)
            lines.append(f"Record Header: {entry.get('record_header', '')}")
            lines.append(f"Location: {entry.get('location', '')}")
            lines.append(f"Height: {entry.get('height', '')}")
            lines.append(f"Skin Color: {entry.get('skin_color', '')}")
            lines.append(f"Hair Color: {entry.get('hair_color', '')}")
            lines.append(f"Hair Texture: {entry.get('hair_texture', '')}")
            lines.append(f"Beard: {entry.get('beard', '')}")
            lines.append(f"Mustache: {entry.get('mustache', '')}")
            lines.append(f"Assignatura: {entry.get('assignatura', '')}")
            lines.append(f"Reservista: {entry.get('reservista', '')}")
            lines.append(f"Eyes: {entry.get('eyes', '')}")
            lines.append(f"Mouth: {entry.get('mouth', '')}")
            lines.append(f"Face: {entry.get('face', '')}")
            lines.append(f"Nose: {entry.get('nose', '')}")
            lines.append(f"Marks: {entry.get('marks', '')}")
            officials = entry.get("officials", [])
            if officials:
                officials_str = "; ".join(
                    [f"{o.get('position', '')}: {o.get('signature', '')}" for o in officials]
                )
            else:
                officials_str = ""
            lines.append(f"Officials: {officials_str}")
            lines.append(f"Father: {entry.get('father', '')}")
            lines.append(f"Mother: {entry.get('mother', '')}")
            lines.append(f"Birth Date: {entry.get('birth_date', '')}")
            lines.append(f"Birth Place: {entry.get('birth_place', '')}")
            lines.append(f"Municipality: {entry.get('municipality', '')}")
            lines.append(f"Civil Status: {entry.get('civil_status', '')}")
            lines.append(f"Vaccinated: {entry.get('vaccinated', '')}")
            lines.append(f"Can Read: {entry.get('can_read', '')}")
            lines.append(f"Can Write: {entry.get('can_write', '')}")
            lines.append(f"Can Count: {entry.get('can_count', '')}")
            lines.append(f"Swimming: {entry.get('swimming', '')}")
            lines.append(f"Cyclist: {entry.get('cyclist', '')}")
            lines.append(f"Motorcyclist: {entry.get('motorcyclist', '')}")
            lines.append(f"Driver: {entry.get('driver', '')}")
            lines.append(f"Chauffeur: {entry.get('chauffeur', '')}")
            lines.append(f"Telegraphist: {entry.get('telegraphist', '')}")
            lines.append(f"Telephonist: {entry.get('telephonist', '')}")
            lines.append(f"Residence: {entry.get('residence', '')}")
            lines.append(f"Observations: {entry.get('observations', '')}")
            lines.append("\n" + "=" * 40 + "\n")
        return lines
