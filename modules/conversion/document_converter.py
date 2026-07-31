"""Document conversion utilities for DOCX and TXT output formats."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document

if TYPE_CHECKING:
    from docx.document import Document as _DocxDocument

from modules.conversion.base import BaseConverter, resolve_field

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Field-spec driven helpers for simple schemas.  Each "field spec" is a tuple
# of (label, dict_key, default_value).  For nested dicts a dotted key like
# "address.street" is supported (one level only).
# ---------------------------------------------------------------------------


def _as_dict(value: Any) -> dict:
    """Return *value* when it is a dict, else an empty dict.

    Nested objects are schema-valid as ``null``, so callers must not assume a
    mapping is present.
    """
    return value if isinstance(value, dict) else {}


def _star_count(value: Any) -> int:
    """Coerce a star count to a non-negative int, tolerating strings/nulls."""
    try:
        stars = int(value)
    except (TypeError, ValueError):
        return 0
    return max(stars, 0)


def _format_geography(geography: Any) -> str:
    """Format a v3.0 geography object as 'city, country' (modern preferred)."""
    geo = _as_dict(geography)
    city = geo.get("city_modern") or geo.get("city_original")
    country = geo.get("country_modern") or geo.get("country_original")
    return ", ".join(str(part) for part in [city, country] if part)


def _format_lifespan(lifespan: Any) -> str:
    """Format a v3.0 lifespan object as 'birth - death'."""
    span = _as_dict(lifespan)
    birth = span.get("birth_year")
    death = span.get("death_year")
    if birth is None and death is None:
        return ""
    start = birth if birth is not None else "?"
    end = death if death is not None else "?"
    return f"{start} - {end}"


def _association_lines(associations: Any) -> list[str]:
    """Render a v3.0 associations list as display strings."""
    if not isinstance(associations, list):
        return []
    lines: list[str] = []
    for assoc in associations:
        if not isinstance(assoc, dict):
            continue
        label = assoc.get("entity_label_modern") or assoc.get("entity_label_original")
        parts = [str(p) for p in [assoc.get("entity_type"), label] if p]
        text = " - ".join(parts)
        relationship = assoc.get("relationship") or assoc.get("role")
        if relationship:
            text = f"{text} ({relationship})" if text else str(relationship)
        if text:
            lines.append(text)
    return lines


def _event_lines(events: Any) -> list[str]:
    """Render a v3.0 places ``events`` list as display strings."""
    if not isinstance(events, list):
        return []
    lines: list[str] = []
    for event in events:
        if not isinstance(event, dict):
            continue
        head = " ".join(
            str(p)
            for p in [event.get("event_type"), event.get("year")]
            if p is not None and p != ""
        )
        description = event.get("description")
        text = (
            f"{head}: {description}"
            if head and description
            else (head or str(description or ""))
        )
        if text:
            lines.append(text)
    return lines


def _utensil_lines(utensils: Any) -> list[str]:
    """Render a recipe ``utensils_equipment`` list with its v3.0 ratings."""
    if not isinstance(utensils, list):
        return []
    lines: list[str] = []
    for utensil in utensils:
        if not isinstance(utensil, dict):
            continue
        name = (
            utensil.get("utensil_modern_english")
            or utensil.get("utensil_original")
            or ""
        )
        ratings = ", ".join(
            f"{label}: {val}"
            for label, val in [
                ("Specialization", utensil.get("utensil_specialization_rating_1_7")),
                ("Modernity", utensil.get("utensil_modernity_rating_1_7")),
            ]
            if val is not None
        )
        text = f"{name} [{ratings}]" if ratings else str(name)
        if text.strip():
            lines.append(text)
    return lines


def _contributor_lines(contributors: Any) -> list[str]:
    """Render a v3.0 works ``contributors`` list as display strings."""
    if not isinstance(contributors, list):
        return []
    lines: list[str] = []
    for contrib in contributors:
        if not isinstance(contrib, dict):
            continue
        name = contrib.get("name_original") or contrib.get("name_modern_english")
        role = contrib.get("role")
        if name and role:
            lines.append(f"{name} ({role})")
        elif name or role:
            lines.append(str(name or role))
    return lines


def _fields_to_docx(
    entries: list[Any],
    document: _DocxDocument,
    header_fn: Any,
    fields: list[tuple],
    *,
    page_break: bool = True,
) -> None:
    """Render *entries* into *document* using a flat field list."""
    for entry in entries:
        document.add_heading(header_fn(entry), level=1)
        for label, key, default in fields:
            value = resolve_field(entry, key, default)
            document.add_paragraph(f"{label}: {value}")
        if page_break:
            document.add_page_break()


def _fields_to_txt(
    entries: list[Any],
    header_fn: Any,
    fields: list[tuple],
    *,
    separator: str = "\n" + "=" * 40 + "\n",
) -> list[str]:
    """Render *entries* into lines using a flat field list."""
    lines: list[str] = []
    for entry in entries:
        lines.append(header_fn(entry))
        for label, key, default in fields:
            value = resolve_field(entry, key, default)
            lines.append(f"{label}: {value}")
        lines.append(separator)
    return lines


class DocumentConverter(BaseConverter):
    """
    Converts JSON-extracted data to DOCX or TXT documents.

    Inherits from BaseConverter for shared entry extraction and utility methods.
    """

    def convert(self, json_file: Path, output_file: Path) -> None:
        """
        Convert JSON to output format based on file extension.

        :param json_file: Input JSON file path
        :param output_file: Output file path (.docx or .txt)
        """
        suffix = output_file.suffix.lower()
        if suffix == ".docx":
            self.convert_to_docx(json_file, output_file)
        elif suffix == ".txt":
            self.convert_to_txt(json_file, output_file)
        else:
            logger.warning(f"Unsupported output format: {suffix}")

    def convert_to_docx(self, json_file: Path, output_file: Path) -> None:
        """Convert JSON entries to a DOCX document."""
        entries = self.get_entries(json_file)
        document: _DocxDocument = Document()
        document.add_heading(json_file.stem, 0)
        converters = {
            "structuredsummaries": self._convert_structured_summaries_to_docx,
            "bibliographicentries": self._convert_bibliographic_entries_to_docx,
            "historicaladdressbookentries": (
                self._convert_historicaladdressbookentries_to_docx
            ),
            "brazilianmilitaryrecords": (
                self._convert_brazilianoccupationrecords_to_docx
            ),
            "culinarypersonsentries": self._convert_culinary_persons_to_docx,
            "culinaryplacesentries": self._convert_culinary_places_to_docx,
            "culinaryworksentries": self._convert_culinary_works_to_docx,
            "culinaryentitiesentries": self._convert_culinary_entities_to_docx,
            "historicalrecipesentriesproduction": (
                self._convert_historical_recipes_production_to_docx
            ),
            "historicalrecipesentriesproductionv3": (
                self._convert_historical_recipes_production_to_docx
            ),
            "michelinguideslight": self._convert_michelin_guides_light_to_docx,
            "cookbookmetadataentries": self._convert_cookbook_metadata_to_docx,
        }
        converter = self.get_converter(converters)
        try:
            if converter:
                converter(entries, document)
            else:
                for entry in entries:
                    document.add_paragraph(self.safe_str(entry))
        except Exception as e:
            logger.warning(
                f"DOCX converter for schema '{self.schema_name}' failed ({e}); "
                "falling back to plain paragraphs."
            )
            for entry in entries:
                document.add_paragraph(self.safe_str(entry))
        try:
            document.save(str(output_file))
            logger.info(f"DOCX file generated at {output_file}")
        except Exception as e:
            logger.error(f"Error saving DOCX file {output_file}: {e}")

    def convert_to_txt(self, json_file: Path, output_file: Path) -> None:
        """Convert JSON entries to a plain text file."""
        entries = self.get_entries(json_file)

        if not entries:
            logger.warning(f"No valid entries found in {json_file.name}")
            with output_file.open("w", encoding="utf-8") as f:
                f.write(f"No valid entries found in {json_file.name}\n")
            return

        converters = {
            "structuredsummaries": self._convert_structured_summaries_to_txt,
            "bibliographicentries": self._convert_bibliographic_entries_to_txt,
            "historicaladdressbookentries": (
                self._convert_historicaladdressbookentries_to_txt
            ),
            "brazilianmilitaryrecords": (
                self._convert_brazilianoccupationrecords_to_txt
            ),
            "culinarypersonsentries": self._convert_culinary_persons_to_txt,
            "culinaryplacesentries": self._convert_culinary_places_to_txt,
            "culinaryworksentries": self._convert_culinary_works_to_txt,
            "culinaryentitiesentries": self._convert_culinary_entities_to_txt,
            "historicalrecipesentriesproduction": (
                self._convert_historical_recipes_production_to_txt
            ),
            "historicalrecipesentriesproductionv3": (
                self._convert_historical_recipes_production_to_txt
            ),
            "michelinguideslight": self._convert_michelin_guides_light_to_txt,
            "cookbookmetadataentries": self._convert_cookbook_metadata_to_txt,
        }

        converter = self.get_converter(converters)
        try:
            if converter:
                lines = converter(entries)
            else:
                lines = [self.safe_str(entry) for entry in entries]
        except Exception as e:
            logger.warning(
                f"TXT converter for schema '{self.schema_name}' failed ({e}); "
                "falling back to plain lines."
            )
            lines = [self.safe_str(entry) for entry in entries]

        try:
            lines = [line for line in lines if line is not None]

            with output_file.open("w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            logger.info(f"TXT file generated at {output_file}")
        except Exception as e:
            logger.error(f"Error writing TXT file {output_file}: {e}")

    # --- Schema-Specific DOCX Converters ---
    def _convert_structured_summaries_to_docx(
        self, entries: list, document: _DocxDocument
    ) -> None:
        """
        Converts structured summaries entries to a DOCX document (schema v4.0).

        For each entry, writes the page number in bold followed by bullet-pointed
        summaries and any references. Reads the schema keys ``page_number``
        (nested integer), ``bullet_points``, and ``references``.
        """
        list_bullet = document.styles["List Bullet"]
        reference_set: set[str] = set()
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            page_info = entry.get("page_number") or {}
            page = (
                page_info.get("page_number_integer")
                if isinstance(page_info, dict)
                else None
            )
            bullet_points = entry.get("bullet_points")
            references = entry.get("references")

            # Add a paragraph with bold text for the page number.
            p_page = document.add_paragraph()
            run_page = p_page.add_run(
                f"Page {page}" if page is not None else "Page Unknown"
            )
            run_page.bold = True

            # Add each bullet point as a separate bullet item.
            if bullet_points and isinstance(bullet_points, list):
                for bp in bullet_points:
                    if bp:
                        p_bp = document.add_paragraph(style=list_bullet)
                        p_bp.add_run(str(bp))
            else:
                document.add_paragraph("No bullet points available.")

            # Add references if available, and accumulate for the final section.
            if references and isinstance(references, list):
                formatted_refs = ", ".join(str(ref) for ref in references if ref)
                if formatted_refs:
                    p_ref = document.add_paragraph(style=list_bullet)
                    p_ref.add_run(f"References: {formatted_refs}")
                for ref in references:
                    if ref:
                        reference_set.add(str(ref))

            # Add an empty paragraph for spacing.
            document.add_paragraph("")

        # If any references were found, add a consolidated final section.
        if reference_set:
            document.add_page_break()
            document.add_heading("References", level=1)
            for ref in sorted(reference_set):
                document.add_paragraph(ref, style=list_bullet)

    def _convert_bibliographic_entries_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Converts bibliographic entries to a DOCX document (schema v4.4).

        Mirrors the TXT twin: entry-level ``full_title``, ``short_title``,
        ``main_author``, ``institutional_main_author``, ``short_note``,
        ``library_abbreviation``, ``volumes_overview``, ``volume_numbers``;
        edition-level ``year``, ``edition_number``, ``publication_locations``,
        ``contributors``, ``edition_category``, ``language``,
        ``translated_from``.
        """
        list_bullet = document.styles["List Bullet"]
        for entry in entries:
            if not isinstance(entry, dict):
                continue

            document.add_heading(
                self.safe_str(entry.get("full_title") or "Unknown Title"), level=1
            )
            document.add_paragraph(
                f"Short Title: {self.safe_str(entry.get('short_title') or '')}"
            )
            document.add_paragraph(
                f"Main Author: {self.safe_str(entry.get('main_author') or 'Anonymous')}"
            )

            inst = entry.get("institutional_main_author")
            if inst is not None:
                document.add_paragraph(f"Institutional Author: {self.safe_str(inst)}")

            short_note = entry.get("short_note")
            if short_note:
                document.add_paragraph(f"Note: {self.safe_str(short_note)}")

            library = entry.get("library_abbreviation")
            if library:
                document.add_paragraph(f"Library: {self.safe_str(library)}")

            volumes_overview = entry.get("volumes_overview")
            if volumes_overview:
                document.add_paragraph(f"Volumes: {self.safe_str(volumes_overview)}")

            volume_numbers = self.join_list(entry.get("volume_numbers"))
            if volume_numbers:
                document.add_paragraph(f"Volume Numbers: {volume_numbers}")

            document.add_heading("Edition Information", level=2)
            editions = entry.get("edition_info") or []
            if not isinstance(editions, list) or not editions:
                document.add_paragraph(
                    "No edition information available.", style=list_bullet
                )
                document.add_page_break()
                continue

            for edition in editions:
                if not isinstance(edition, dict):
                    continue

                pub_locations = edition.get("publication_locations") or []
                places = [
                    loc.get("modern_place") or loc.get("original_place")
                    for loc in pub_locations
                    if isinstance(loc, dict)
                    and (loc.get("modern_place") or loc.get("original_place"))
                ]
                location_str = ", ".join(str(p) for p in places) or "Unknown"

                contributors = edition.get("contributors") or []
                contributor_strs = [
                    f"{self.safe_str(c.get('name'))} ({self.safe_str(c.get('role'))})"
                    for c in contributors
                    if isinstance(c, dict)
                ]
                contributors_str = (
                    ", ".join(contributor_strs) if contributor_strs else "Unknown"
                )

                ed_cat = self.safe_str(edition.get("edition_category") or "")
                edition_text = (
                    f"Year: {self.safe_str(edition.get('year') or 'Unknown')}, "
                    f"Edition:"
                    f" {self.safe_str(edition.get('edition_number') or 'Unknown')}, "
                    f"Location: {location_str}, "
                    f"Contributors: {contributors_str}, "
                    f"Category: {ed_cat}, "
                    f"Language: {self.safe_str(edition.get('language') or '')}, "
                    f"Translated From:"
                    f" {self.safe_str(edition.get('translated_from') or '')}"
                )
                document.add_paragraph(edition_text, style=list_bullet)

            document.add_page_break()

    def _convert_culinary_entities_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Converts unified culinary entities entries to DOCX (schema v3.0)."""
        for entry in entries:
            if not isinstance(entry, dict):
                continue

            entry_type, profile = self._entity_profile(entry)
            names = _as_dict(profile.get("names"))
            title = (
                names.get("original")
                or names.get("modern_english")
                or f"{entry_type} Entry"
            )
            document.add_heading(f"{title} ({entry_type})", level=1)

            timeframe = _as_dict(profile.get("timeframe"))
            geography = _as_dict(profile.get("geography"))

            def add_paragraph(label: str, value: Any) -> None:
                if value not in (None, ""):
                    document.add_paragraph(f"{label}: {value}")

            add_paragraph("Modern Name", names.get("modern_english"))
            add_paragraph("Importance", profile.get("importance"))
            add_paragraph("Summary", profile.get("summary"))
            add_paragraph("Timeframe", timeframe.get("notation"))
            add_paragraph("Timeframe Start", timeframe.get("start_year"))
            add_paragraph("Timeframe End", timeframe.get("end_year"))
            add_paragraph("Primary Location", geography.get("primary_location"))
            add_paragraph("Geographic Context", geography.get("additional_context"))
            add_paragraph("Topical Focus", self.join_list(profile.get("topical_focus")))
            add_paragraph("Languages", self.join_list(profile.get("language_contexts")))

            if entry_type == "Person":
                add_paragraph("Roles", self.join_list(profile.get("roles")))

            elif entry_type == "Place":
                add_paragraph(
                    "Culinary Roles",
                    self.join_list(profile.get("roles_in_culinary_ecosystem")),
                )
                add_paragraph(
                    "Associated Products",
                    self.join_list(profile.get("associated_products")),
                )
                add_paragraph(
                    "Notable Establishments",
                    self.join_list(profile.get("notable_establishments")),
                )
                add_paragraph("Place Notes", profile.get("place_notes"))

            elif entry_type == "Work":
                add_paragraph("Short Title", profile.get("short_title"))
                add_paragraph("Genre", profile.get("genre"))

            document.add_page_break()

    _ENTITY_PROFILE_KEYS = {
        "Person": "person_entry",
        "Place": "place_entry",
        "Work": "work_entry",
    }

    @classmethod
    def _entity_profile(cls, entry: dict) -> tuple[str, dict]:
        """Return ``(entry_type, profile_dict)`` for a culinary entities entry."""
        entry_type = str(entry.get("entry_type") or "Unknown")
        profile_key = cls._ENTITY_PROFILE_KEYS.get(entry_type)
        profile = entry.get(profile_key) if profile_key else None
        return entry_type, _as_dict(profile)

    def _convert_culinary_entities_to_txt(self, entries: list[Any]) -> list[str]:
        """Converts unified culinary entities entries to TXT (schema v3.0)."""
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue

            entry_type, profile = self._entity_profile(entry)
            names = _as_dict(profile.get("names"))
            header = (
                names.get("original")
                or names.get("modern_english")
                or f"{entry_type} Entry"
            )
            lines.append(f"Entry Type: {entry_type}")
            lines.append(f"Name: {header}")

            timeframe = _as_dict(profile.get("timeframe"))
            geography = _as_dict(profile.get("geography"))
            lines.append(f"  Importance: {self.safe_str(profile.get('importance'))}")
            lines.append(f"  Summary: {self.safe_str(profile.get('summary'))}")
            lines.append(f"  Timeframe: {self.safe_str(timeframe.get('notation'))}")
            lines.append(
                f"  Timeframe Start: {self.safe_str(timeframe.get('start_year'))}"
            )
            lines.append(f"  Timeframe End: {self.safe_str(timeframe.get('end_year'))}")
            lines.append(
                f"  Primary Location:"
                f" {self.safe_str(geography.get('primary_location'))}"
            )
            lines.append(
                f"  Geographic Context:"
                f" {self.safe_str(geography.get('additional_context'))}"
            )
            lines.append(
                f"  Topical Focus: {self.join_list(profile.get('topical_focus'))}"
            )
            lines.append(
                f"  Languages: {self.join_list(profile.get('language_contexts'))}"
            )

            if entry_type == "Person":
                lines.append(f"  Roles: {self.join_list(profile.get('roles'))}")

            elif entry_type == "Place":
                culinary_roles = self.join_list(
                    profile.get("roles_in_culinary_ecosystem")
                )
                lines.append(f"  Culinary Roles: {culinary_roles}")
                lines.append(
                    f"  Associated Products:"
                    f" {self.join_list(profile.get('associated_products'))}"
                )
                lines.append(
                    f"  Notable Establishments:"
                    f" {self.join_list(profile.get('notable_establishments'))}"
                )
                lines.append(
                    f"  Place Notes: {self.safe_str(profile.get('place_notes'))}"
                )

            elif entry_type == "Work":
                lines.append(
                    f"  Short Title: {self.safe_str(profile.get('short_title'))}"
                )
                lines.append(f"  Genre: {self.safe_str(profile.get('genre'))}")

            lines.append("")

        return lines

    # --- Shared field specs for simple schemas ---

    _ADDRESSBOOK_FIELDS: list[tuple] = [
        ("Address", "address.street", "Unknown"),
        ("Street Number", "address.street_number", ""),
        ("Honorific", "honorific", ""),
        ("Additional Notes", "additional_notes", ""),
    ]

    @staticmethod
    def _addressbook_header(entry: dict) -> str:
        last = entry.get("last_name", "Unknown")
        first = entry.get("first_name", "Unknown")
        occupation = entry.get("occupation", "Unknown")
        header = f"{last}, {first} - {occupation}"
        section = entry.get("section")
        if section:
            header += f" (Section: {section})"
        return header

    def _convert_historicaladdressbookentries_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        _fields_to_docx(
            entries, document, self._addressbook_header, self._ADDRESSBOOK_FIELDS
        )

    _BRAZILIAN_RECORDS_FIELDS: list[tuple] = [
        ("Record Header", "record_header", ""),
        ("Location", "location", ""),
        ("Height", "height", ""),
        ("Skin Color", "skin_color", ""),
        ("Hair Color", "hair_color", ""),
        ("Hair Texture", "hair_texture", ""),
        ("Beard", "beard", ""),
        ("Mustache", "mustache", ""),
        ("Assignatura", "assignatura", ""),
        ("Reservista", "reservista", ""),
        ("Eyes", "eyes", ""),
        ("Mouth", "mouth", ""),
        ("Face", "face", ""),
        ("Nose", "nose", ""),
        ("Marks", "marks", ""),
        ("Father", "father", ""),
        ("Mother", "mother", ""),
        ("Birth Date", "birth_date", ""),
        ("Birth Place", "birth_place", ""),
        ("Municipality", "municipality", ""),
        ("Civil Status", "civil_status", ""),
        ("Vaccinated", "vaccinated", ""),
        ("Can Read", "can_read", ""),
        ("Can Write", "can_write", ""),
        ("Can Count", "can_count", ""),
        ("Swimming", "swimming", ""),
        ("Cyclist", "cyclist", ""),
        ("Motorcyclist", "motorcyclist", ""),
        ("Driver", "driver", ""),
        ("Chauffeur", "chauffeur", ""),
        ("Telegraphist", "telegraphist", ""),
        ("Telephonist", "telephonist", ""),
        ("Residence", "residence", ""),
        ("Observations", "observations", ""),
    ]

    @staticmethod
    def _brazilian_header(entry: dict) -> str:
        surname = entry.get("surname", "")
        first = entry.get("first_name", "")
        profession = entry.get("profession", "")
        return f"{surname}, {first} - {profession}"

    def _convert_brazilianoccupationrecords_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        for entry in entries:
            document.add_heading(self._brazilian_header(entry), level=1)
            for label, key, default in self._BRAZILIAN_RECORDS_FIELDS:
                document.add_paragraph(f"{label}: {resolve_field(entry, key, default)}")
            document.add_paragraph(f"Officials: {self._format_officials(entry)}")
            document.add_page_break()

    # --- Schema-Specific TXT Converters ---
    def _convert_structured_summaries_to_txt(self, entries: list[Any]) -> list[str]:
        """Converts structured summaries entries to TXT format (schema v4.0)."""
        lines: list[str] = []
        reference_set: set[str] = set()
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            page_info = entry.get("page_number") or {}
            page = (
                page_info.get("page_number_integer")
                if isinstance(page_info, dict)
                else None
            )
            bullet_points = entry.get("bullet_points", [])
            references = entry.get("references", [])
            lines.append(f"Page {page}" if page is not None else "Page Unknown")
            # List the bullet points.
            if bullet_points and isinstance(bullet_points, list):
                for bp in bullet_points:
                    if bp:
                        lines.append(f" - {bp}")
            else:
                lines.append("No bullet points available.")
            # Append a bullet point for references and collect for final section.
            if references and isinstance(references, list):
                formatted_refs = ", ".join(str(ref) for ref in references if ref)
                if formatted_refs:
                    lines.append(f" - References: {formatted_refs}")
                for ref in references:
                    if ref:
                        reference_set.add(str(ref))
            lines.append("")
        if reference_set:
            lines.append("References:")
            for ref in sorted(reference_set):
                lines.append(f" - {ref}")
        return lines

    def _convert_bibliographic_entries_to_txt(self, entries: list[Any]) -> list[str]:
        """Converts bibliographic entries to TXT format (schema v4.4).

        Reads the current schema keys: entry-level ``full_title``,
        ``short_title``, ``main_author``, ``institutional_main_author``,
        ``short_note``, ``library_abbreviation``, ``volumes_overview``,
        ``volume_numbers``; edition-level ``year``, ``edition_number``,
        ``publication_locations`` (modern/original place), ``contributors``
        (name/role), ``edition_category``, ``language``, ``translated_from``.
        """
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue

            lines.append(
                f"Full Title: {self.safe_str(entry.get('full_title', 'Unknown Title'))}"
            )
            lines.append(f"Short Title: {self.safe_str(entry.get('short_title', ''))}")
            lines.append(
                f"Main Author: {self.safe_str(entry.get('main_author', 'Anonymous'))}"
            )
            inst = entry.get("institutional_main_author")
            if inst is not None:
                lines.append(f"Institutional Author: {self.safe_str(inst)}")

            short_note = entry.get("short_note")
            if short_note:
                lines.append(f"Note: {self.safe_str(short_note)}")

            library = entry.get("library_abbreviation")
            if library:
                lines.append(f"Library: {self.safe_str(library)}")

            volumes_overview = entry.get("volumes_overview")
            if volumes_overview:
                lines.append(f"Volumes: {self.safe_str(volumes_overview)}")

            volume_numbers = self.join_list(entry.get("volume_numbers"))
            if volume_numbers:
                lines.append(f"Volume Numbers: {volume_numbers}")

            lines.append("Edition Information:")
            editions = entry.get("edition_info", [])
            if not isinstance(editions, list):
                editions = []

            for edition in editions:
                if not isinstance(edition, dict):
                    continue

                pub_locations = edition.get("publication_locations") or []
                places = [
                    loc.get("modern_place") or loc.get("original_place")
                    for loc in pub_locations
                    if isinstance(loc, dict)
                    and (loc.get("modern_place") or loc.get("original_place"))
                ]
                location_str = ", ".join(str(p) for p in places) or "Unknown"

                contributors = edition.get("contributors") or []
                contributor_strs = [
                    f"{self.safe_str(c.get('name'))} ({self.safe_str(c.get('role'))})"
                    for c in contributors
                    if isinstance(c, dict)
                ]
                contributors_str = (
                    ", ".join(contributor_strs) if contributor_strs else "Unknown"
                )

                ed_year = self.safe_str(edition.get("year", "Unknown"))
                ed_num = self.safe_str(edition.get("edition_number", "Unknown"))
                ed_cat = self.safe_str(edition.get("edition_category", ""))
                ed_lang = self.safe_str(edition.get("language", ""))
                ed_trans = self.safe_str(edition.get("translated_from", ""))
                edition_text = (
                    f"Year: {ed_year}, "
                    f"Edition: {ed_num}, "
                    f"Location: {location_str}, "
                    f"Contributors: {contributors_str}, "
                    f"Category: {ed_cat}, "
                    f"Language: {ed_lang}, "
                    f"Translated From: {ed_trans}"
                )
                lines.append(f" - {edition_text}")

            lines.append("\n" + "=" * 40 + "\n")

        return lines

    def _convert_historicaladdressbookentries_to_txt(
        self, entries: list[Any]
    ) -> list[str]:
        return _fields_to_txt(
            entries, self._addressbook_header, self._ADDRESSBOOK_FIELDS, separator=""
        )

    def _convert_brazilianoccupationrecords_to_txt(
        self, entries: list[Any]
    ) -> list[str]:
        lines: list[str] = []
        for entry in entries:
            lines.append(self._brazilian_header(entry))
            for label, key, default in self._BRAZILIAN_RECORDS_FIELDS:
                lines.append(f"{label}: {resolve_field(entry, key, default)}")
            lines.append(f"Officials: {self._format_officials(entry)}")
            lines.append("\n" + "=" * 40 + "\n")
        return lines

    # --- Culinary Schemas DOCX Converters ---
    def _convert_culinary_persons_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Converts culinary persons entries to DOCX format (schema v3.0)."""
        entries = self._normalize_entries(entries)
        list_bullet = document.styles["List Bullet"]
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            names = _as_dict(entry.get("names"))
            name = names.get("original") or names.get("modern_english") or "Unknown"
            document.add_heading(str(name), level=1)

            modern_name = names.get("modern_english")
            if modern_name and modern_name != name:
                document.add_paragraph(f"Modern Name: {modern_name}")

            gender = entry.get("gender")
            if gender:
                document.add_paragraph(f"Gender: {gender}")

            roles = self.join_list(entry.get("roles"))
            if roles:
                document.add_paragraph(f"Roles: {roles}")

            importance = entry.get("historical_importance")
            if importance is not None:
                document.add_paragraph(f"Historical Importance: {importance}/7")

            period_str = self._format_period(entry)
            if period_str:
                document.add_paragraph(f"Timeframe: {period_str}")

            lifespan_str = _format_lifespan(entry.get("lifespan"))
            if lifespan_str:
                document.add_paragraph(f"Lifespan: {lifespan_str}")

            geography_str = _format_geography(entry.get("geography"))
            if geography_str:
                document.add_paragraph(f"Geography: {geography_str}")

            associations = _association_lines(entry.get("associations"))
            if associations:
                document.add_heading("Associations", level=2)
                for assoc in associations:
                    document.add_paragraph(assoc, style=list_bullet)

            short_notes = entry.get("short_notes")
            if short_notes:
                document.add_heading("Notes", level=2)
                document.add_paragraph(str(short_notes))

            document.add_page_break()

    def _convert_culinary_places_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Converts culinary places entries to DOCX format (schema v3.0)."""
        entries = self._normalize_entries(entries)
        list_bullet = document.styles["List Bullet"]
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            names = _as_dict(entry.get("names"))
            name = names.get("original") or names.get("modern_english") or "Unknown"
            document.add_heading(str(name), level=1)

            modern_name = names.get("modern_english")
            if modern_name and modern_name != name:
                document.add_paragraph(f"Modern Name: {modern_name}")

            place_type = entry.get("place_type")
            if place_type:
                document.add_paragraph(f"Type: {place_type}")

            geography_str = _format_geography(entry.get("geography"))
            if geography_str:
                document.add_paragraph(f"Geography: {geography_str}")

            importance = entry.get("historical_importance")
            if importance is not None:
                document.add_paragraph(f"Historical Importance: {importance}/7")

            period_str = self._format_period(entry)
            if period_str:
                document.add_paragraph(f"Timeframe: {period_str}")

            roles = self.join_list(entry.get("roles_in_culinary_ecosystem"))
            if roles:
                document.add_paragraph(f"Roles: {roles}")

            events = _event_lines(entry.get("events"))
            if events:
                document.add_heading("Events", level=2)
                for event in events:
                    document.add_paragraph(event, style=list_bullet)

            associations = _association_lines(entry.get("associations"))
            if associations:
                document.add_heading("Associations", level=2)
                for assoc in associations:
                    document.add_paragraph(assoc, style=list_bullet)

            short_notes = entry.get("short_notes")
            if short_notes:
                document.add_heading("Notes", level=2)
                document.add_paragraph(str(short_notes))

            document.add_page_break()

    def _convert_culinary_works_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Converts culinary works entries to DOCX format (schema v3.0)."""
        entries = self._normalize_entries(entries)
        list_bullet = document.styles["List Bullet"]
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            titles = _as_dict(entry.get("titles"))
            title = titles.get("original") or titles.get("modern_english") or "Unknown"
            document.add_heading(str(title), level=1)

            modern_title = titles.get("modern_english")
            if modern_title and modern_title != title:
                document.add_paragraph(f"Modern Title: {modern_title}")

            short_title = titles.get("short")
            if short_title:
                document.add_paragraph(f"Short Title: {short_title}")

            genre = entry.get("genre")
            if genre:
                document.add_paragraph(f"Genre: {genre}")

            importance = entry.get("historical_importance")
            if importance is not None:
                document.add_paragraph(f"Historical Importance: {importance}/7")

            culinary_focus = self.join_list(entry.get("culinary_focus"))
            if culinary_focus:
                document.add_paragraph(f"Culinary Focus: {culinary_focus}")

            languages = self.join_list(entry.get("languages"))
            if languages:
                document.add_paragraph(f"Languages: {languages}")

            edition_years = self.join_list(entry.get("edition_years"))
            if edition_years:
                document.add_paragraph(f"Edition Years: {edition_years}")

            period_str = self._format_period(entry)
            if period_str:
                document.add_paragraph(f"Timeframe: {period_str}")

            geography_str = _format_geography(entry.get("geography"))
            if geography_str:
                document.add_paragraph(f"Geography: {geography_str}")

            contributors = _contributor_lines(entry.get("contributors"))
            if contributors:
                document.add_heading("Contributors", level=2)
                for contrib in contributors:
                    document.add_paragraph(contrib, style=list_bullet)

            associations = _association_lines(entry.get("associations"))
            if associations:
                document.add_heading("Associations", level=2)
                for assoc in associations:
                    document.add_paragraph(assoc, style=list_bullet)

            short_notes = entry.get("short_notes")
            if short_notes:
                document.add_heading("Notes", level=2)
                document.add_paragraph(str(short_notes))

            document.add_page_break()

    # --- Culinary Schemas TXT Converters ---
    def _convert_culinary_persons_to_txt(self, entries: list[Any]) -> list[str]:
        """Converts culinary persons entries to TXT format (schema v3.0)."""
        entries = self._normalize_entries(entries)
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            names = _as_dict(entry.get("names"))
            name = names.get("original") or names.get("modern_english") or "Unknown"
            lines.append(str(name))

            modern_name = names.get("modern_english")
            if modern_name and modern_name != name:
                lines.append(f"Modern Name: {modern_name}")

            gender = entry.get("gender")
            if gender:
                lines.append(f"Gender: {gender}")

            roles = self.join_list(entry.get("roles"))
            if roles:
                lines.append(f"Roles: {roles}")

            importance = entry.get("historical_importance")
            if importance is not None:
                lines.append(f"Historical Importance: {importance}/7")

            period_str = self._format_period(entry)
            if period_str:
                lines.append(f"Timeframe: {period_str}")

            lifespan_str = _format_lifespan(entry.get("lifespan"))
            if lifespan_str:
                lines.append(f"Lifespan: {lifespan_str}")

            geography_str = _format_geography(entry.get("geography"))
            if geography_str:
                lines.append(f"Geography: {geography_str}")

            associations = _association_lines(entry.get("associations"))
            if associations:
                lines.append("Associations:")
                lines.extend(f" - {assoc}" for assoc in associations)

            short_notes = entry.get("short_notes")
            if short_notes:
                lines.append(f"Notes: {short_notes}")

            lines.append("\n" + "=" * 40 + "\n")
        return lines

    def _convert_culinary_places_to_txt(self, entries: list[Any]) -> list[str]:
        """Converts culinary places entries to TXT format (schema v3.0)."""
        entries = self._normalize_entries(entries)
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            names = _as_dict(entry.get("names"))
            name = names.get("original") or names.get("modern_english") or "Unknown"
            lines.append(str(name))

            modern_name = names.get("modern_english")
            if modern_name and modern_name != name:
                lines.append(f"Modern Name: {modern_name}")

            place_type = entry.get("place_type")
            if place_type:
                lines.append(f"Type: {place_type}")

            geography_str = _format_geography(entry.get("geography"))
            if geography_str:
                lines.append(f"Geography: {geography_str}")

            importance = entry.get("historical_importance")
            if importance is not None:
                lines.append(f"Historical Importance: {importance}/7")

            period_str = self._format_period(entry)
            if period_str:
                lines.append(f"Timeframe: {period_str}")

            roles = self.join_list(entry.get("roles_in_culinary_ecosystem"))
            if roles:
                lines.append(f"Roles: {roles}")

            events = _event_lines(entry.get("events"))
            if events:
                lines.append("Events:")
                lines.extend(f" - {event}" for event in events)

            associations = _association_lines(entry.get("associations"))
            if associations:
                lines.append("Associations:")
                lines.extend(f" - {assoc}" for assoc in associations)

            short_notes = entry.get("short_notes")
            if short_notes:
                lines.append(f"Notes: {short_notes}")

            lines.append("\n" + "=" * 40 + "\n")
        return lines

    def _convert_culinary_works_to_txt(self, entries: list[Any]) -> list[str]:
        """Converts culinary works entries to TXT format (schema v3.0)."""
        entries = self._normalize_entries(entries)
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            titles = _as_dict(entry.get("titles"))
            title = titles.get("original") or titles.get("modern_english") or "Unknown"
            lines.append(str(title))

            modern_title = titles.get("modern_english")
            if modern_title and modern_title != title:
                lines.append(f"Modern Title: {modern_title}")

            short_title = titles.get("short")
            if short_title:
                lines.append(f"Short Title: {short_title}")

            genre = entry.get("genre")
            if genre:
                lines.append(f"Genre: {genre}")

            importance = entry.get("historical_importance")
            if importance is not None:
                lines.append(f"Historical Importance: {importance}/7")

            culinary_focus = self.join_list(entry.get("culinary_focus"))
            if culinary_focus:
                lines.append(f"Culinary Focus: {culinary_focus}")

            languages = self.join_list(entry.get("languages"))
            if languages:
                lines.append(f"Languages: {languages}")

            edition_years = self.join_list(entry.get("edition_years"))
            if edition_years:
                lines.append(f"Edition Years: {edition_years}")

            period_str = self._format_period(entry)
            if period_str:
                lines.append(f"Timeframe: {period_str}")

            geography_str = _format_geography(entry.get("geography"))
            if geography_str:
                lines.append(f"Geography: {geography_str}")

            contributors = _contributor_lines(entry.get("contributors"))
            if contributors:
                lines.append("Contributors:")
                lines.extend(f" - {contrib}" for contrib in contributors)

            associations = _association_lines(entry.get("associations"))
            if associations:
                lines.append("Associations:")
                lines.extend(f" - {assoc}" for assoc in associations)

            short_notes = entry.get("short_notes")
            if short_notes:
                lines.append(f"Notes: {short_notes}")

            lines.append("\n" + "=" * 40 + "\n")
        return lines

    # --- Michelin Guides Light Converters (schema 3.4-light) ---

    def _convert_michelin_guides_light_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Convert MichelinGuidesLight entries to DOCX format (schema 3.4-light)."""
        entries = self._normalize_entries(entries)
        for entry in entries:
            if not isinstance(entry, dict):
                continue

            name = entry.get("establishment_name", "Unknown Establishment")
            awards = entry.get("awards", {}) or {}
            stars = _star_count(awards.get("stars"))
            star_display = "⭐" * stars if stars else ""
            document.add_heading(f"{name} {star_display}".strip(), level=1)

            location = entry.get("location", {}) or {}
            address = entry.get("address", {}) or {}
            location_parts = [
                p
                for p in [
                    location.get("neighbourhood_or_area"),
                    location.get("city_or_town"),
                ]
                if p
            ]
            if location_parts:
                document.add_paragraph(f"Location: {', '.join(location_parts)}")
            address_parts = [
                str(p)
                for p in [
                    address.get("street"),
                    address.get("house_number"),
                    address.get("postal_code"),
                ]
                if p
            ]
            if address_parts:
                document.add_paragraph(f"Address: {' '.join(address_parts)}")

            award_bits: list[str] = []
            if awards.get("bib_gourmand"):
                award_bits.append("Bib Gourmand")
            if awards.get("michelin_plate"):
                award_bits.append("Michelin Plate")
            if awards.get("pleasant_marker"):
                award_bits.append("Pleasant")
            if awards.get("hotel_class"):
                award_bits.append(f"Hotel class {awards.get('hotel_class')}")
            if awards.get("restaurant_class"):
                award_bits.append(f"Restaurant class {awards.get('restaurant_class')}")
            if award_bits:
                document.add_paragraph(f"Awards: {', '.join(award_bits)}")

            cuisine = entry.get("cuisine", {}) or {}
            origin = self.join_list(cuisine.get("cuisine_origin"))
            if origin:
                document.add_paragraph(f"Cuisine: {origin}")
            style = self.join_list(cuisine.get("culinary_style"))
            if style:
                document.add_paragraph(f"Style: {style}")
            specialties = self.join_list(cuisine.get("specialties"))
            if specialties:
                document.add_paragraph(f"Specialties: {specialties}")

            pricing = entry.get("pricing", {}) or {}
            currency = pricing.get("currency") or ""
            menu_min = pricing.get("menu_price_min")
            menu_max = pricing.get("menu_price_max")
            if menu_min or menu_max:
                document.add_paragraph(
                    f"Menu Price: {currency} {menu_min or '?'} - {menu_max or '?'}"
                )
            alc_min = pricing.get("a_la_carte_price_min")
            alc_max = pricing.get("a_la_carte_price_max")
            if alc_min or alc_max:
                document.add_paragraph(
                    f"À la carte: {currency} {alc_min or '?'} - {alc_max or '?'}"
                )

            note = entry.get("inspector_note")
            if note:
                document.add_paragraph(f"Inspector Note: {note}")

            document.add_page_break()

    def _convert_michelin_guides_light_to_txt(self, entries: list[Any]) -> list[str]:
        """Convert MichelinGuidesLight entries to TXT format (schema 3.4-light)."""
        entries = self._normalize_entries(entries)
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue

            name = entry.get("establishment_name", "Unknown Establishment")
            awards = entry.get("awards", {}) or {}
            stars = _star_count(awards.get("stars"))
            star_display = "*" * stars if stars else "No stars"

            lines.append(f"{'=' * 60}")
            lines.append(f"{name}")
            lines.append(f"Stars: {star_display}")
            lines.append(f"{'=' * 60}")

            location = entry.get("location", {}) or {}
            address = entry.get("address", {}) or {}
            loc_parts = [
                p
                for p in [
                    location.get("neighbourhood_or_area"),
                    location.get("city_or_town"),
                ]
                if p
            ]
            if loc_parts:
                lines.append(f"Location: {', '.join(loc_parts)}")
            addr_parts = [
                str(p)
                for p in [
                    address.get("street"),
                    address.get("house_number"),
                    address.get("postal_code"),
                ]
                if p
            ]
            if addr_parts:
                lines.append(f"Address: {' '.join(addr_parts)}")

            award_bits: list[str] = []
            if awards.get("bib_gourmand"):
                award_bits.append("Bib Gourmand")
            if awards.get("michelin_plate"):
                award_bits.append("Michelin Plate")
            if awards.get("pleasant_marker"):
                award_bits.append("Pleasant")
            if awards.get("hotel_class"):
                award_bits.append(f"Hotel class {awards.get('hotel_class')}")
            if awards.get("restaurant_class"):
                award_bits.append(f"Restaurant class {awards.get('restaurant_class')}")
            if award_bits:
                lines.append(f"Awards: {', '.join(award_bits)}")

            cuisine = entry.get("cuisine", {}) or {}
            origin = self.join_list(cuisine.get("cuisine_origin"))
            if origin:
                lines.append(f"Cuisine: {origin}")
            style = self.join_list(cuisine.get("culinary_style"))
            if style:
                lines.append(f"Style: {style}")
            specialties = self.join_list(cuisine.get("specialties"))
            if specialties:
                lines.append(f"Specialties: {specialties}")

            pricing = entry.get("pricing", {}) or {}
            currency = pricing.get("currency") or ""
            menu_min = pricing.get("menu_price_min")
            menu_max = pricing.get("menu_price_max")
            if menu_min or menu_max:
                lines.append(f"Menu: {currency} {menu_min or '?'} - {menu_max or '?'}")
            alc_min = pricing.get("a_la_carte_price_min")
            alc_max = pricing.get("a_la_carte_price_max")
            if alc_min or alc_max:
                lines.append(
                    f"À la carte: {currency} {alc_min or '?'} - {alc_max or '?'}"
                )

            note = entry.get("inspector_note")
            if note:
                lines.append(f"Inspector Note: {note}")

            lines.append("")

        return lines

    def _convert_cookbook_metadata_to_txt(self, entries: list[Any]) -> list[str]:
        """Convert cookbook metadata entries to the required plain text format."""
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            lines.append(f"title: {self.safe_str(entry.get('title') or 'unknown')}")
            lines.append(f"author: {self.safe_str(entry.get('author') or 'anonymous')}")
            lines.append(f"year: {self.safe_str(entry.get('year') or 'unknown')}")
            lines.append(f"edition: {self.safe_str(entry.get('edition') or 'unknown')}")
            lines.append(f"content: {self.safe_str(entry.get('content') or '')}")
            lines.append(f"notes: {self.safe_str(entry.get('notes') or '')}")
            lines.append(f"library: {self.safe_str(entry.get('library') or 'unknown')}")
            lines.append(
                f"digitizer: {self.safe_str(entry.get('digitizer') or 'unknown')}"
            )
            lines.append(f"misc: {self.safe_str(entry.get('misc') or '')}")
            lines.append("\n" + "=" * 40 + "\n")
        return lines

    # --- CookbookMetadataEntries DOCX Converter ---

    _COOKBOOK_METADATA_FIELDS: list[tuple] = [
        ("Author", "author", "anonymous"),
        ("Year", "year", "unknown"),
        ("Edition", "edition", "unknown"),
        ("Content", "content", ""),
        ("Notes", "notes", ""),
        ("Library", "library", "unknown"),
        ("Digitizer", "digitizer", "unknown"),
        ("Misc", "misc", ""),
    ]

    @staticmethod
    def _cookbook_metadata_header(entry: dict) -> str:
        return entry.get("title") or "Unknown Title"

    def _convert_cookbook_metadata_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Converts cookbook metadata entries to DOCX format."""
        _fields_to_docx(
            entries,
            document,
            self._cookbook_metadata_header,
            self._COOKBOOK_METADATA_FIELDS,
        )

    # --- HistoricalRecipesEntriesProduction Converters ---

    def _convert_historical_recipes_production_to_docx(
        self, entries: list[Any], document: _DocxDocument
    ) -> None:
        """Converts HistoricalRecipesEntriesProduction entries to DOCX (schema v3.0)."""
        entries = self._normalize_entries(entries)
        list_bullet = document.styles["List Bullet"]
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            title = entry.get("title_original") or "Unknown Recipe"
            document.add_heading(str(title), level=1)

            modern_title = entry.get("title_modern_english")
            if modern_title and modern_title != title:
                document.add_paragraph(f"Modern Title: {modern_title}")

            recipe_type = entry.get("recipe_type")
            if recipe_type:
                document.add_paragraph(f"Type: {recipe_type}")

            # Timing/yield — stored in a nested timing_yield object
            timing_yield = entry.get("timing_yield", {}) or {}
            yield_str = timing_yield.get("yield_original") or ""
            prep_str = timing_yield.get("preparation_time_original") or ""
            cook_str = timing_yield.get("cooking_time_original") or ""
            if yield_str:
                document.add_paragraph(f"Yield: {yield_str}")
            if prep_str:
                document.add_paragraph(f"Preparation Time: {prep_str}")
            if cook_str:
                document.add_paragraph(f"Cooking Time: {cook_str}")

            # Ingredients with per-ingredient ratings
            ingredients = entry.get("ingredients", [])
            if ingredients:
                document.add_heading("Ingredients", level=2)
                for ing in ingredients:
                    name = (
                        ing.get("name_modern_english") or ing.get("name_original") or ""
                    )
                    qty = ing.get("quantity_original") or ""
                    ing_text = f"{name} ({qty})" if qty else name
                    luxury = ing.get("ingredient_luxury_signal_rating_1_7")
                    trade = ing.get("ingredient_trade_distance_rating_1_7")
                    novelty = ing.get("ingredient_novelty_rating_1_7")
                    ratings = ", ".join(
                        f"{label}: {val}"
                        for label, val in [
                            ("Luxury", luxury),
                            ("Trade dist.", trade),
                            ("Novelty", novelty),
                        ]
                        if val is not None
                    )
                    if ratings:
                        ing_text += f" [{ratings}]"
                    document.add_paragraph(ing_text, style=list_bullet)

            # Cooking methods with per-method complexity rating
            methods = entry.get("cooking_methods", [])
            if methods:
                document.add_heading("Cooking Methods", level=2)
                for m in methods:
                    method_name = (
                        m.get("method_modern_english") or m.get("method_original") or ""
                    )
                    complexity = m.get("method_complexity_rating_1_7")
                    method_text = method_name
                    if complexity is not None:
                        method_text += f" [Complexity: {complexity}]"
                    document.add_paragraph(method_text, style=list_bullet)

            # Utensils and equipment with per-utensil ratings
            utensils = _utensil_lines(entry.get("utensils_equipment"))
            if utensils:
                document.add_heading("Utensils and Equipment", level=2)
                for utensil in utensils:
                    document.add_paragraph(utensil, style=list_bullet)

            # Culinary style
            culinary_style = entry.get("culinary_style", {}) or {}
            modernity = culinary_style.get("modernity_rating_1_7")
            if modernity is not None:
                document.add_paragraph(f"Modernity Rating: {modernity}/7")
            innovation = culinary_style.get("innovation_markers_observed", [])
            if innovation:
                document.add_paragraph(f"Innovation Markers: {', '.join(innovation)}")
            archaism = culinary_style.get("archaism_markers_observed", [])
            if archaism:
                document.add_paragraph(f"Archaism Markers: {', '.join(archaism)}")

            # Original recipe text
            recipe_text = entry.get("recipe_text_original")
            if recipe_text:
                document.add_heading("Original Recipe Text", level=2)
                document.add_paragraph(recipe_text)

            # Modern translation
            recipe_text_modern = entry.get("recipe_text_modern_english")
            if recipe_text_modern and recipe_text_modern != recipe_text:
                document.add_heading("Modern English Translation", level=2)
                document.add_paragraph(recipe_text_modern)

            document.add_page_break()

    def _convert_historical_recipes_production_to_txt(
        self, entries: list[Any]
    ) -> list[str]:
        """Converts HistoricalRecipesEntriesProduction entries to TXT (schema v3.0)."""
        entries = self._normalize_entries(entries)
        lines: list[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            title = entry.get("title_original") or "Unknown Recipe"
            lines.append(str(title))

            modern_title = entry.get("title_modern_english")
            if modern_title and modern_title != title:
                lines.append(f"Modern Title: {modern_title}")

            recipe_type = entry.get("recipe_type")
            if recipe_type:
                lines.append(f"Type: {recipe_type}")

            # Timing/yield — stored in a nested timing_yield object
            timing_yield = entry.get("timing_yield", {}) or {}
            yield_str = timing_yield.get("yield_original") or ""
            prep_str = timing_yield.get("preparation_time_original") or ""
            cook_str = timing_yield.get("cooking_time_original") or ""
            if yield_str:
                lines.append(f"Yield: {yield_str}")
            if prep_str:
                lines.append(f"Preparation Time: {prep_str}")
            if cook_str:
                lines.append(f"Cooking Time: {cook_str}")

            # Ingredients with per-ingredient ratings
            ingredients = entry.get("ingredients", [])
            if ingredients:
                lines.append("Ingredients:")
                for ing in ingredients:
                    name = (
                        ing.get("name_modern_english") or ing.get("name_original") or ""
                    )
                    qty = ing.get("quantity_original") or ""
                    ing_text = f" - {name} ({qty})" if qty else f" - {name}"
                    luxury = ing.get("ingredient_luxury_signal_rating_1_7")
                    trade = ing.get("ingredient_trade_distance_rating_1_7")
                    novelty = ing.get("ingredient_novelty_rating_1_7")
                    ratings = ", ".join(
                        f"{label}: {val}"
                        for label, val in [
                            ("Luxury", luxury),
                            ("Trade dist.", trade),
                            ("Novelty", novelty),
                        ]
                        if val is not None
                    )
                    if ratings:
                        ing_text += f" [{ratings}]"
                    lines.append(ing_text)

            # Cooking methods with complexity rating
            methods = entry.get("cooking_methods", [])
            if methods:
                method_parts: list[str] = []
                for m in methods:
                    method_name = (
                        m.get("method_modern_english") or m.get("method_original") or ""
                    )
                    complexity = m.get("method_complexity_rating_1_7")
                    part = method_name
                    if complexity is not None:
                        part += f" [Complexity: {complexity}]"
                    method_parts.append(part)
                lines.append(f"Cooking Methods: {', '.join(method_parts)}")

            # Utensils and equipment with per-utensil ratings
            utensils = _utensil_lines(entry.get("utensils_equipment"))
            if utensils:
                lines.append("Utensils and Equipment:")
                lines.extend(f" - {utensil}" for utensil in utensils)

            # Culinary style
            culinary_style = entry.get("culinary_style", {}) or {}
            modernity = culinary_style.get("modernity_rating_1_7")
            if modernity is not None:
                lines.append(f"Modernity Rating: {modernity}/7")
            innovation = culinary_style.get("innovation_markers_observed", [])
            if innovation:
                lines.append(f"Innovation Markers: {', '.join(innovation)}")
            archaism = culinary_style.get("archaism_markers_observed", [])
            if archaism:
                lines.append(f"Archaism Markers: {', '.join(archaism)}")

            # Original recipe text
            recipe_text = entry.get("recipe_text_original")
            if recipe_text:
                lines.append("Original Recipe Text:")
                lines.append(recipe_text)

            # Modern translation
            recipe_text_modern = entry.get("recipe_text_modern_english")
            if recipe_text_modern and recipe_text_modern != recipe_text:
                lines.append("Modern English Translation:")
                lines.append(recipe_text_modern)

            lines.append("\n" + "=" * 40 + "\n")
        return lines
