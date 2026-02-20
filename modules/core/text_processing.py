# modules/core/text_processing.py

"""Document conversion utilities for DOCX and TXT output formats."""

import logging
from pathlib import Path
from typing import Any, List

from docx import Document

from modules.core.converter_base import BaseConverter, resolve_field

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Field-spec driven helpers for simple schemas.  Each "field spec" is a tuple
# of (label, dict_key, default_value).  For nested dicts a dotted key like
# "address.street" is supported (one level only).
# ---------------------------------------------------------------------------

def _fields_to_docx(
    entries: List[Any],
    document: "Document",
    header_fn: Any,
    fields: List[tuple],
    *,
    page_break: bool = True,
) -> None:
    """Render *entries* into *document* using a flat field list."""
    for entry in entries:
        document.add_heading(header_fn(entry), level=1)
        for label, key, default in fields:
            value = resolve_field(entry, key, default)
            document.add_paragraph(f"{label}: {value}")
        if page_break:
            document.add_page_break()


def _fields_to_txt(
    entries: List[Any],
    header_fn: Any,
    fields: List[tuple],
    *,
    separator: str = "\n" + "=" * 40 + "\n",
) -> List[str]:
    """Render *entries* into lines using a flat field list."""
    lines: List[str] = []
    for entry in entries:
        lines.append(header_fn(entry))
        for label, key, default in fields:
            value = resolve_field(entry, key, default)
            lines.append(f"{label}: {value}")
        lines.append(separator)
    return lines


class DocumentConverter(BaseConverter):
    """
    Converts JSON-extracted data to DOCX or TXT documents.
    
    Inherits from BaseConverter for shared entry extraction and utility methods.
    """
    
    def convert(self, json_file: Path, output_file: Path) -> None:
        """
        Convert JSON to output format based on file extension.
        
        :param json_file: Input JSON file path
        :param output_file: Output file path (.docx or .txt)
        """
        suffix = output_file.suffix.lower()
        if suffix == ".docx":
            self.convert_to_docx(json_file, output_file)
        elif suffix == ".txt":
            self.convert_to_txt(json_file, output_file)
        else:
            logger.warning(f"Unsupported output format: {suffix}")
    
    def convert_to_docx(self, json_file: Path, output_file: Path) -> None:
        """Convert JSON entries to a DOCX document."""
        entries = self.get_entries(json_file)
        document: Document = Document()
        document.add_heading(json_file.stem, 0)
        converters = {
            "structuredsummaries": self._convert_structured_summaries_to_docx,
            "bibliographicentries": self._convert_bibliographic_entries_to_docx,
            "historicaladdressbookentries": self._convert_historicaladdressbookentries_to_docx,
            "brazilianoccupationrecords": self._convert_brazilianoccupationrecords_to_docx,
            "brazilianmilitaryrecords": self._convert_brazilianoccupationrecords_to_docx,
            "culinarypersonsentries": self._convert_culinary_persons_to_docx,
            "culinaryplacesentries": self._convert_culinary_places_to_docx,
            "culinaryworksentries": self._convert_culinary_works_to_docx,
            "culinaryentitiesentries": self._convert_culinary_entities_to_docx,
            "historicalrecipesentries": self._convert_historical_recipes_to_docx,
            "michelinguides": self._convert_michelin_guides_to_docx
        }
        converter = self.get_converter(converters)
        if converter:
            converter(entries, document)
        else:
            for entry in entries:
                document.add_paragraph(str(entry))
        try:
            document.save(output_file)
            logger.info(f"DOCX file generated at {output_file}")
        except Exception as e:
            logger.error(f"Error saving DOCX file {output_file}: {e}")

    def convert_to_txt(self, json_file: Path, output_file: Path) -> None:
        """Convert JSON entries to a plain text file."""
        entries = self.get_entries(json_file)

        if not entries:
            logger.warning(f"No valid entries found in {json_file.name}")
            with output_file.open("w", encoding="utf-8") as f:
                f.write(f"No valid entries found in {json_file.name}\n")
            return

        converters = {
            "structuredsummaries": self._convert_structured_summaries_to_txt,
            "bibliographicentries": self._convert_bibliographic_entries_to_txt,
            "historicaladdressbookentries": self._convert_historicaladdressbookentries_to_txt,
            "brazilianmilitaryrecords": self._convert_brazilianoccupationrecords_to_txt,
            "brazilianoccupationrecords": self._convert_brazilianoccupationrecords_to_txt,
            "culinarypersonsentries": self._convert_culinary_persons_to_txt,
            "culinaryplacesentries": self._convert_culinary_places_to_txt,
            "culinaryworksentries": self._convert_culinary_works_to_txt,
            "culinaryentitiesentries": self._convert_culinary_entities_to_txt,
            "historicalrecipesentries": self._convert_historical_recipes_to_txt,
            "michelinguides": self._convert_michelin_guides_to_txt,
            "cookbookmetadataentries": self._convert_cookbook_metadata_to_txt
        }

        try:
            converter = self.get_converter(converters)
            if converter:
                lines = converter(entries)
            else:
                lines = [self.safe_str(entry) for entry in entries]

            lines = [line for line in lines if line is not None]

            with output_file.open("w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            logger.info(f"TXT file generated at {output_file}")
        except Exception as e:
            logger.error(f"Error writing TXT file {output_file}: {e}")

    # --- Schema-Specific DOCX Converters ---
    def _convert_structured_summaries_to_docx(self, entries: list,
                                              document: "Document") -> None:
        """
        Converts structured summaries entries to a DOCX document.
        For each entry, writes the page number in bold (as simple text) followed by bullet-pointed summaries.
        The keywords are formatted in italic (without asterisks).
        """
        literature_set = set()
        for entry in entries:
            page = entry.get("page")
            bullet_points = entry.get("bullet_points")
            keywords = entry.get("keywords")
            literature = entry.get("literature")

            # Instead of adding a heading, add a paragraph with bold text for the page number.
            p_page = document.add_paragraph()
            run_page = p_page.add_run(
                f"Page {page}" if page is not None else "Page Unknown")
            run_page.bold = True

            # For keywords: add a bullet point with italic text, without asterisks.
            if keywords and isinstance(keywords, list):
                formatted_keywords = ", ".join(kw for kw in keywords if kw)
                p_keyword = document.add_paragraph(style='List Bullet')
                p_keyword.add_run("Keywords: ")
                run_keywords = p_keyword.add_run(formatted_keywords)
                run_keywords.italic = True

            # Add each bullet point as a separate bullet item.
            if bullet_points and isinstance(bullet_points, list):
                for bp in bullet_points:
                    if bp:
                        p_bp = document.add_paragraph(style='List Bullet')
                        p_bp.add_run(str(bp))
            else:
                document.add_paragraph("No bullet points available.")

            # Add literature references if available.
            if literature and isinstance(literature, list):
                formatted_refs = ", ".join(
                    str(ref) for ref in literature if ref)
                p_ref = document.add_paragraph(style='List Bullet')
                p_ref.add_run(f"References: {formatted_refs}")

            # Accumulate literature for the consolidated literature section.
            if literature and isinstance(literature, list):
                for lit in literature:
                    if lit:
                        literature_set.add(lit)

            # Add an empty paragraph for spacing.
            document.add_paragraph("")

        # If any literature was found, add a final section.
        if literature_set:
            document.add_page_break()
            document.add_heading("Literature", level=1)
            for lit in sorted(literature_set):
                document.add_paragraph(str(lit), style='List Bullet')

    def _convert_bibliographic_entries_to_docx(self, entries: List[Any],
                                               document: Document) -> None:
        for entry in entries:
            # Extract primary entry data
            full_title = entry.get("full_title", "Unknown")
            short_title = entry.get("short_title", "")
            culinary_focus = entry.get("culinary_focus", [])
            edition_info = entry.get("edition_info", [])
            total_editions = entry.get("total_editions", "")

            first_edition = None
            if isinstance(edition_info, list) and edition_info:
                first_edition = edition_info[0] if isinstance(edition_info[0], dict) else None

            book_format = entry.get("format")
            if not book_format and isinstance(first_edition, dict):
                book_format = first_edition.get("format")

            pages = entry.get("pages")
            if pages in (None, "") and isinstance(first_edition, dict):
                pages = first_edition.get("pages")

            # Add entry header and basic information
            document.add_heading(full_title, level=1)
            document.add_paragraph(f"Short Title: {short_title}")

            # Add culinary focus areas
            if culinary_focus and isinstance(culinary_focus, list):
                document.add_paragraph(
                    f"Culinary Focus: {', '.join(culinary_focus)}")
            else:
                document.add_paragraph("Culinary Focus: Unknown")

            # Add format and page information
            document.add_paragraph(
                f"Format: {book_format if book_format else 'Unknown'}")
            document.add_paragraph(
                f"Pages: {pages if pages is not None else 'Unknown'}")
            document.add_paragraph(
                f"Total Editions: {total_editions if total_editions is not None else 'Unknown'}")

            # Add edition information
            document.add_heading("Edition Information", level=2)
            if edition_info and isinstance(edition_info, list):
                for edition in edition_info:
                    if not isinstance(edition, dict):
                        continue

                    # Extract edition data
                    edition_number = edition.get("edition_number", "Unknown")
                    year = edition.get("year", "Unknown")

                    # Extract location data
                    city = "Unknown"
                    country = "Unknown"
                    publication_locations = edition.get("publication_locations")
                    if isinstance(publication_locations, list) and publication_locations:
                        first_loc = publication_locations[0] if isinstance(publication_locations[0], dict) else {}
                        if isinstance(first_loc, dict):
                            city = (
                                first_loc.get("modern_place")
                                or first_loc.get("original_place")
                                or "Unknown"
                            )
                            country = (
                                first_loc.get("modern_region")
                                or first_loc.get("original_region")
                                or "Unknown"
                            )
                    else:
                        # Backward compatibility
                        location = edition.get("location", {})
                        if isinstance(location, dict):
                            country = location.get("country") or "Unknown"
                            city = location.get("city") or "Unknown"

                    # Get contributors
                    contributors = edition.get("contributors", [])
                    contributors_str = "Unknown"
                    if isinstance(contributors, list) and contributors:
                        formatted_contributors = []
                        for c in contributors:
                            if isinstance(c, dict):
                                name = c.get("name")
                                role = c.get("role")
                                if name and role:
                                    formatted_contributors.append(f"{name} ({role})")
                                elif name:
                                    formatted_contributors.append(str(name))
                            elif c is not None:
                                formatted_contributors.append(str(c))
                        formatted_contributors = [x for x in formatted_contributors if x]
                        if formatted_contributors:
                            contributors_str = ", ".join(formatted_contributors)

                    # Get other edition details
                    language = edition.get("language", "")
                    short_note = edition.get("short_note", "")
                    edition_category = edition.get("edition_category", "")

                    # Format edition text
                    edition_text = (
                        f"Edition: {edition_number if edition_number is not None else 'Unknown'}, "
                        f"Year: {year if year is not None else 'Unknown'}, "
                        f"Location: {city}, {country}, "
                        f"Language: {language if language else 'Unknown'}, "
                        f"Contributors: {contributors_str}, "
                        f"Category: {edition_category if edition_category else 'Unknown'}"
                    )

                    if short_note:
                        edition_text += f"\nNote: {short_note}"

                    document.add_paragraph(edition_text, style='List Bullet')
            else:
                document.add_paragraph("No edition information available.",
                                       style='List Bullet')

            document.add_page_break()

    def _convert_culinary_entities_to_docx(self, entries: List[Any], document: Document) -> None:
        profile_keys = {
            "Person": "person_entry",
            "Place": "place_entry",
            "Work": "work_entry"
        }

        for entry in entries:
            if not isinstance(entry, dict):
                continue

            entry_type = entry.get("entry_type", "Unknown")
            profile = entry.get(profile_keys.get(entry_type, ""), {})
            if not isinstance(profile, dict):
                profile = {}

            names = profile.get("names", {}) or {}
            title = names.get("original") or names.get("modern_english") or f"{entry_type} Entry"
            document.add_heading(f"{title} ({entry_type})", level=1)

            timeframe = profile.get("timeframe", {}) or {}
            topical_focus = self.join_list(profile.get("topical_focus"))
            languages = self.join_list(profile.get("language_contexts"))
            associations = self.format_associations(profile.get("associations"), as_list=True)

            def add_paragraph(label: str, value: Any) -> None:
                if value not in (None, ""):
                    document.add_paragraph(f"{label}: {value}")

            add_paragraph("Modern Name", names.get("modern_english"))
            add_paragraph("Summary", profile.get("entity_summary"))
            add_paragraph("Timeframe", timeframe.get("notation"))
            add_paragraph("Timeframe Start", timeframe.get("start_year"))
            add_paragraph("Timeframe End", timeframe.get("end_year"))
            add_paragraph("Topical Focus", topical_focus)
            add_paragraph("Languages", languages)
            add_paragraph("Notes", profile.get("notes"))

            if entry_type == "Person":
                add_paragraph("Gender", profile.get("gender"))
                add_paragraph("Roles", self.join_list(profile.get("roles")))
                add_paragraph("Name Variants", self.format_name_variants(profile.get("name_variants")))
                add_paragraph("Biographical Notes", profile.get("biographical_notes"))

            elif entry_type == "Place":
                add_paragraph("Place Type", profile.get("place_type"))
                add_paragraph("Country (Modern)", profile.get("country_modern"))
                add_paragraph("Culinary Roles", self.join_list(profile.get("roles_in_culinary_ecosystem")))
                add_paragraph("Associated Products", self.join_list(profile.get("associated_products")))
                add_paragraph("Notable Establishments", self.join_list(profile.get("notable_establishments")))
                add_paragraph("Place Notes", profile.get("place_notes"))

            elif entry_type == "Work":
                add_paragraph("Short Title", profile.get("short_title"))
                add_paragraph("Description", profile.get("description"))
                add_paragraph("Genre", profile.get("genre"))
                add_paragraph("Edition Years", self.join_list(profile.get("edition_years")))
                material = profile.get("material_features", {}) or {}
                add_paragraph("Format", material.get("format"))
                add_paragraph("Has Illustrations", material.get("has_illustrations"))
                add_paragraph("Page Count", material.get("page_count"))
                add_paragraph("Material Notes", material.get("notes"))

            if associations:
                document.add_paragraph("Associations:")
                for assoc in associations:
                    document.add_paragraph(assoc, style='List Bullet')

            document.add_page_break()

    def _convert_culinary_entities_to_txt(self, entries: List[Any]) -> List[str]:
        profile_keys = {
            "Person": "person_entry",
            "Place": "place_entry",
            "Work": "work_entry"
        }

        lines: List[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue

            entry_type = entry.get("entry_type", "Unknown")
            profile = entry.get(profile_keys.get(entry_type, ""), {})
            if not isinstance(profile, dict):
                profile = {}

            names = profile.get("names", {}) or {}
            header = names.get("original") or names.get("modern_english") or f"{entry_type} Entry"
            lines.append(f"Entry Type: {entry_type}")
            lines.append(f"Name: {header}")

            timeframe = profile.get("timeframe", {}) or {}
            lines.append(f"  Summary: {self.safe_str(profile.get('entity_summary'))}")
            lines.append(f"  Timeframe: {self.safe_str(timeframe.get('notation'))}")
            lines.append(f"  Timeframe Start: {self.safe_str(timeframe.get('start_year'))}")
            lines.append(f"  Timeframe End: {self.safe_str(timeframe.get('end_year'))}")
            lines.append(f"  Topical Focus: {self.join_list(profile.get('topical_focus'))}")
            lines.append(f"  Languages: {self.join_list(profile.get('language_contexts'))}")
            lines.append(f"  Associations: {self.format_associations(profile.get('associations'))}")
            lines.append(f"  Notes: {self.safe_str(profile.get('notes'))}")

            if entry_type == "Person":
                lines.append(f"  Gender: {self.safe_str(profile.get('gender'))}")
                lines.append(f"  Roles: {self.join_list(profile.get('roles'))}")
                lines.append(f"  Name Variants: {self.format_name_variants(profile.get('name_variants'))}")
                lines.append(f"  Biographical Notes: {self.safe_str(profile.get('biographical_notes'))}")

            elif entry_type == "Place":
                lines.append(f"  Place Type: {self.safe_str(profile.get('place_type'))}")
                lines.append(f"  Country (Modern): {self.safe_str(profile.get('country_modern'))}")
                lines.append(f"  Culinary Roles: {self.join_list(profile.get('roles_in_culinary_ecosystem'))}")
                lines.append(f"  Associated Products: {self.join_list(profile.get('associated_products'))}")
                lines.append(f"  Notable Establishments: {self.join_list(profile.get('notable_establishments'))}")
                lines.append(f"  Place Notes: {self.safe_str(profile.get('place_notes'))}")

            elif entry_type == "Work":
                lines.append(f"  Short Title: {self.safe_str(profile.get('short_title'))}")
                lines.append(f"  Description: {self.safe_str(profile.get('description'))}")
                lines.append(f"  Genre: {self.safe_str(profile.get('genre'))}")
                lines.append(f"  Edition Years: {self.join_list(profile.get('edition_years'))}")
                material = profile.get("material_features", {}) or {}
                lines.append(f"  Format: {self.safe_str(material.get('format'))}")
                lines.append(f"  Has Illustrations: {self.safe_str(material.get('has_illustrations'))}")
                lines.append(f"  Page Count: {self.safe_str(material.get('page_count'))}")
                lines.append(f"  Material Notes: {self.safe_str(material.get('notes'))}")

            lines.append("")

        return lines

    # --- Shared field specs for simple schemas ---

    _ADDRESSBOOK_FIELDS: List[tuple] = [
        ("Address", "address.street", "Unknown"),
        ("Street Number", "address.street_number", "*0*"),
        ("Honorific", "honorific", ""),
        ("Additional Notes", "additional_notes", ""),
    ]

    @staticmethod
    def _addressbook_header(entry: dict) -> str:
        header = f"{entry.get('last_name', 'Unknown')}, {entry.get('first_name', 'Unknown')} - {entry.get('occupation', 'Unknown')}"
        section = entry.get("section")
        if section:
            header += f" (Section: {section})"
        return header

    def _convert_historicaladdressbookentries_to_docx(self, entries: List[Any],
                                                      document: Document) -> None:
        _fields_to_docx(entries, document, self._addressbook_header, self._ADDRESSBOOK_FIELDS)

    _BRAZILIAN_RECORDS_FIELDS: List[tuple] = [
        ("Record Header", "record_header", ""),
        ("Location", "location", ""),
        ("Height", "height", ""),
        ("Skin Color", "skin_color", ""),
        ("Hair Color", "hair_color", ""),
        ("Hair Texture", "hair_texture", ""),
        ("Beard", "beard", ""),
        ("Mustache", "mustache", ""),
        ("Assignatura", "assignatura", ""),
        ("Reservista", "reservista", ""),
        ("Eyes", "eyes", ""),
        ("Mouth", "mouth", ""),
        ("Face", "face", ""),
        ("Nose", "nose", ""),
        ("Marks", "marks", ""),
        ("Father", "father", ""),
        ("Mother", "mother", ""),
        ("Birth Date", "birth_date", ""),
        ("Birth Place", "birth_place", ""),
        ("Municipality", "municipality", ""),
        ("Civil Status", "civil_status", ""),
        ("Vaccinated", "vaccinated", ""),
        ("Can Read", "can_read", ""),
        ("Can Write", "can_write", ""),
        ("Can Count", "can_count", ""),
        ("Swimming", "swimming", ""),
        ("Cyclist", "cyclist", ""),
        ("Motorcyclist", "motorcyclist", ""),
        ("Driver", "driver", ""),
        ("Chauffeur", "chauffeur", ""),
        ("Telegraphist", "telegraphist", ""),
        ("Telephonist", "telephonist", ""),
        ("Residence", "residence", ""),
        ("Observations", "observations", ""),
    ]

    @staticmethod
    def _brazilian_header(entry: dict) -> str:
        return f"{entry.get('surname', '')}, {entry.get('first_name', '')} - {entry.get('profession', '')}"

    @staticmethod
    def _format_officials(entry: dict) -> str:
        officials = entry.get("officials", [])
        if officials:
            return "; ".join(
                f"{o.get('position', '')}: {o.get('signature', '')}" for o in officials
            )
        return ""

    def _convert_brazilianoccupationrecords_to_docx(self, entries: List[Any], document: Document) -> None:
        for entry in entries:
            document.add_heading(self._brazilian_header(entry), level=1)
            for label, key, default in self._BRAZILIAN_RECORDS_FIELDS:
                document.add_paragraph(f"{label}: {resolve_field(entry, key, default)}")
            document.add_paragraph(f"Officials: {self._format_officials(entry)}")
            document.add_page_break()

    # --- Schema-Specific TXT Converters ---
    def _convert_structured_summaries_to_txt(
        self, entries: List[Any]
    ) -> List[str]:
        lines: List[str] = []
        literature_set = set()
        for entry in entries:
            page = entry.get("page", "Unknown")
            bullet_points = entry.get("bullet_points", [])
            keywords = entry.get("keywords", [])
            literature = entry.get("literature", [])
            lines.append(f"Page {page}")
            # Add keywords bullet point.
            if keywords and isinstance(keywords, list):
                formatted_keywords = ", ".join(f"*{kw}*" for kw in keywords if kw)
                lines.append(f" - Keywords: {formatted_keywords}")
            # List the bullet points.
            if bullet_points and isinstance(bullet_points, list):
                for bp in bullet_points:
                    if bp:
                        lines.append(f" - {bp}")
            else:
                lines.append("No bullet points available.")
            # Append a bullet point for literature references.
            if literature and isinstance(literature, list):
                formatted_refs = ", ".join(str(ref) for ref in literature if ref)
                lines.append(f" - References: {formatted_refs}")
            # Collect literature for the final consolidated section.
            if literature and isinstance(literature, list):
                for lit in literature:
                    if lit:
                        literature_set.add(lit)
            lines.append("")
        if literature_set:
            lines.append("Literature:")
            for lit in sorted(literature_set):
                lines.append(f" - {lit}")
        return lines

    def _convert_bibliographic_entries_to_txt(
        self, entries: List[Any]
    ) -> List[str]:
        lines: List[str] = []
        for entry in entries:
            if entry is None:
                continue  # Skip None entries

            lines.append(
                f"Full Title: {self.safe_str(entry.get('full_title', 'Unknown Title'))}")
            lines.append(
                f"Short Title: {self.safe_str(entry.get('short_title', ''))}")
            lines.append(
                f"Bibliography Number: {self.safe_str(entry.get('bibliography_number', ''))}")

            authors = entry.get("authors", [])
            if authors is None:
                authors = []
            # Filter out None values in authors list
            authors = [author for author in authors if author is not None]
            lines.append(
                f"Authors: {', '.join(authors) if authors else 'Anonymous'}")

            roles = entry.get("roles", [])
            if roles is None:
                roles = []
            # Filter out None values in roles list
            roles = [role for role in roles if role is not None]
            lines.append(f"Roles: {', '.join(roles)}")

            culinary_focus = entry.get("culinary_focus", [])
            if culinary_focus is None:
                culinary_focus = []
            # Filter out None values in culinary_focus list
            culinary_focus = [focus for focus in culinary_focus if
                              focus is not None]
            lines.append(f"Culinary Focus: {', '.join(culinary_focus)}")

            lines.append(f"Format: {self.safe_str(entry.get('format', ''))}")
            lines.append(f"Pages: {self.safe_str(entry.get('pages', ''))}")
            lines.append(
                f"Total Editions: {self.safe_str(entry.get('total_editions', ''))}")

            lines.append("Edition Information:")
            editions = entry.get("edition_info", [])
            if editions is None:
                editions = []

            for edition in editions:
                if edition is None:
                    continue  # Skip None editions

                # Safely get location info with null checks
                location = edition.get("location", {}) or {}
                city = self.safe_str(location.get("city", "Unknown"))
                country = self.safe_str(location.get("country", "Unknown"))

                # Safely get roles with null checks
                ed_roles = edition.get("roles", []) or []
                ed_roles = [role for role in ed_roles if role is not None]

                edition_text = (
                    f"Year: {self.safe_str(edition.get('year', 'Unknown'))}, "
                    f"Edition: {self.safe_str(edition.get('edition_number', 'Unknown'))}, "
                    f"Location: {city}, {country}, "
                    f"Roles: {', '.join(ed_roles)}, "
                    f"Note: {self.safe_str(edition.get('short_note', ''))}, "
                    f"Category: {self.safe_str(edition.get('edition_category', ''))}, "
                    f"Language: {self.safe_str(edition.get('language', ''))}, "
                    f"Orig Lang: {self.safe_str(edition.get('original_language', ''))}, "
                    f"Translated From: {self.safe_str(edition.get('translated_from', ''))}"
                )
                lines.append(f" - {edition_text}")

            lines.append("\n" + "=" * 40 + "\n")

        return lines

    def _convert_historicaladdressbookentries_to_txt(
        self, entries: List[Any]
    ) -> List[str]:
        return _fields_to_txt(entries, self._addressbook_header, self._ADDRESSBOOK_FIELDS, separator="")

    def _convert_brazilianoccupationrecords_to_txt(self, entries: List[Any]) -> List[str]:
        lines: List[str] = []
        for entry in entries:
            lines.append(self._brazilian_header(entry))
            for label, key, default in self._BRAZILIAN_RECORDS_FIELDS:
                lines.append(f"{label}: {resolve_field(entry, key, default)}")
            lines.append(f"Officials: {self._format_officials(entry)}")
            lines.append("\n" + "=" * 40 + "\n")
        return lines

    # --- Culinary Schemas DOCX Converters ---
    def _convert_culinary_persons_to_docx(self, entries: List[Any], document: Document) -> None:
        """Converts culinary persons entries to DOCX format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        for entry in entries:
            name = entry.get("canonical_name_original", "Unknown")
            document.add_heading(name, level=1)
            
            modern_name = entry.get("canonical_name_modern_english")
            if modern_name and modern_name != name:
                document.add_paragraph(f"Modern Name: {modern_name}")
            
            gender = entry.get("gender")
            if gender:
                document.add_paragraph(f"Gender: {gender}")
            
            roles = entry.get("roles", [])
            if roles:
                document.add_paragraph(f"Roles: {', '.join(roles)}")
            
            period = entry.get("period", {})
            if period:
                period_str = f"{period.get('start_year', 'Unknown')} - {period.get('end_year', 'Unknown')}"
                if period.get("notation"):
                    period_str += f" ({period['notation']})"
                document.add_paragraph(f"Period: {period_str}")
            
            associated_works = entry.get("associated_works", [])
            if associated_works:
                document.add_heading("Associated Works", level=2)
                for work in associated_works:
                    title = work.get("title_original", "")
                    role = work.get("role", "")
                    document.add_paragraph(f"{title} ({role})", style='List Bullet')
            
            associated_places = entry.get("associated_places", [])
            if associated_places:
                document.add_heading("Associated Places", level=2)
                for place in associated_places:
                    place_name = place.get("place_original", "")
                    assoc_type = place.get("association_type", "")
                    document.add_paragraph(f"{place_name} - {assoc_type}", style='List Bullet')
            
            notes = entry.get("notes")
            if notes:
                document.add_heading("Notes", level=2)
                document.add_paragraph(notes)
            
            document.add_page_break()

    def _convert_culinary_places_to_docx(self, entries: List[Any], document: Document) -> None:
        """Converts culinary places entries to DOCX format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        for entry in entries:
            name = entry.get("name_original", "Unknown")
            document.add_heading(name, level=1)
            
            modern_name = entry.get("name_modern_english")
            if modern_name and modern_name != name:
                document.add_paragraph(f"Modern Name: {modern_name}")
            
            place_type = entry.get("place_type")
            country = entry.get("country_modern")
            if place_type:
                document.add_paragraph(f"Type: {place_type}")
            if country:
                document.add_paragraph(f"Country: {country}")
            
            period = entry.get("period", {})
            if period:
                period_str = f"{period.get('start_year', 'Unknown')} - {period.get('end_year', 'Unknown')}"
                if period.get("notation"):
                    period_str += f" ({period['notation']})"
                document.add_paragraph(f"Period: {period_str}")
            
            roles = entry.get("roles_in_culinary_ecosystem", [])
            if roles:
                document.add_paragraph(f"Roles: {', '.join(roles)}")
            
            products = entry.get("associated_products", [])
            if products:
                document.add_heading("Associated Products", level=2)
                for product in products:
                    document.add_paragraph(product, style='List Bullet')
            
            establishments = entry.get("notable_establishments", [])
            if establishments:
                document.add_heading("Notable Establishments", level=2)
                for est in establishments:
                    document.add_paragraph(est, style='List Bullet')
            
            notes = entry.get("notes")
            if notes:
                document.add_heading("Notes", level=2)
                document.add_paragraph(notes)
            
            document.add_page_break()

    def _convert_culinary_works_to_docx(self, entries: List[Any], document: Document) -> None:
        """Converts culinary works entries to DOCX format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        for entry in entries:
            title = entry.get("title_original", "Unknown")
            document.add_heading(title, level=1)
            
            modern_title = entry.get("title_modern_english")
            if modern_title and modern_title != title:
                document.add_paragraph(f"Modern Title: {modern_title}")
            
            short_title = entry.get("short_title")
            if short_title:
                document.add_paragraph(f"Short Title: {short_title}")
            
            genre = entry.get("genre")
            if genre:
                document.add_paragraph(f"Genre: {genre}")
            
            description = entry.get("description")
            if description:
                document.add_paragraph(f"Description: {description}")
            
            culinary_focus = entry.get("culinary_focus", [])
            if culinary_focus:
                document.add_paragraph(f"Culinary Focus: {', '.join(culinary_focus)}")
            
            languages = entry.get("languages", [])
            if languages:
                document.add_paragraph(f"Languages: {', '.join(languages)}")
            
            edition_years = entry.get("edition_years", [])
            if edition_years:
                years_str = ", ".join([str(y) for y in edition_years if y is not None])
                document.add_paragraph(f"Edition Years: {years_str}")
            
            contributors = entry.get("contributors", [])
            if contributors:
                document.add_heading("Contributors", level=2)
                for contrib in contributors:
                    name = contrib.get("name_original", "")
                    role = contrib.get("role", "")
                    document.add_paragraph(f"{name} ({role})", style='List Bullet')
            
            pub_places = entry.get("publication_places", [])
            if pub_places:
                document.add_heading("Publication Places", level=2)
                for place in pub_places:
                    place_name = place.get("name_original", "")
                    document.add_paragraph(place_name, style='List Bullet')
            
            notes = entry.get("notes")
            if notes:
                document.add_heading("Notes", level=2)
                document.add_paragraph(notes)
            
            document.add_page_break()

    def _convert_historical_recipes_to_docx(self, entries: List[Any], document: Document) -> None:
        """Converts historical recipes entries to DOCX format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        for entry in entries:
            title = entry.get("title_original", "Unknown Recipe")
            document.add_heading(title, level=1)
            
            modern_title = entry.get("title_modern_english")
            if modern_title and modern_title != title:
                document.add_paragraph(f"Modern Title: {modern_title}")
            
            recipe_type = entry.get("recipe_type")
            if recipe_type:
                document.add_paragraph(f"Type: {recipe_type}")
            
            # Yield, prep time, cook time
            yields = entry.get("yield", [])
            if yields and len(yields) > 0:
                y = yields[0]
                val = y.get("value_modern_english")
                unit = y.get("unit_modern_english")
                if val and unit:
                    document.add_paragraph(f"Yield: {val} {unit}")
            
            prep_times = entry.get("preparation_time", [])
            if prep_times and len(prep_times) > 0:
                t = prep_times[0]
                val = t.get("value_modern_english")
                unit = t.get("unit_modern_english")
                if val and unit:
                    document.add_paragraph(f"Preparation Time: {val} {unit}")
            
            cook_times = entry.get("cooking_time", [])
            if cook_times and len(cook_times) > 0:
                t = cook_times[0]
                val = t.get("value_modern_english")
                unit = t.get("unit_modern_english")
                if val and unit:
                    document.add_paragraph(f"Cooking Time: {val} {unit}")
            
            # Ingredients
            ingredients = entry.get("ingredients", [])
            if ingredients:
                document.add_heading("Ingredients", level=2)
                for ing in ingredients:
                    name = ing.get("name_modern_english") or ing.get("name_original") or ""
                    qty = ing.get("quantity_original") or ""
                    prep = ing.get("preparation_note_modern_english") or ""
                    ing_text = name
                    if qty:
                        ing_text += f" - {qty}"
                    if prep:
                        ing_text += f" ({prep})"
                    document.add_paragraph(ing_text, style='List Bullet')
            
            # Cooking methods
            methods = entry.get("cooking_methods", [])
            if methods:
                document.add_heading("Cooking Methods", level=2)
                method_names = [m.get("method_modern_english") or m.get("method_original") or "" for m in methods]
                document.add_paragraph(", ".join(method_names))
            
            # Original recipe text
            recipe_text = entry.get("recipe_text_original")
            if recipe_text:
                document.add_heading("Original Recipe Text", level=2)
                document.add_paragraph(recipe_text)
            
            # Modern translation
            recipe_text_modern = entry.get("recipe_text_modern_english")
            if recipe_text_modern and recipe_text_modern != recipe_text:
                document.add_heading("Modern English Translation", level=2)
                document.add_paragraph(recipe_text_modern)
            
            document.add_page_break()

    # --- Culinary Schemas TXT Converters ---
    def _convert_culinary_persons_to_txt(self, entries: List[Any]) -> List[str]:
        """Converts culinary persons entries to TXT format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        lines: List[str] = []
        for entry in entries:
            name = entry.get("canonical_name_original", "Unknown")
            lines.append(name)
            
            modern_name = entry.get("canonical_name_modern_english")
            if modern_name and modern_name != name:
                lines.append(f"Modern Name: {modern_name}")
            
            gender = entry.get("gender")
            if gender:
                lines.append(f"Gender: {gender}")
            
            roles = entry.get("roles", [])
            if roles:
                lines.append(f"Roles: {', '.join(roles)}")
            
            period = entry.get("period", {})
            if period:
                period_str = f"{period.get('start_year', 'Unknown')} - {period.get('end_year', 'Unknown')}"
                if period.get("notation"):
                    period_str += f" ({period['notation']})"
                lines.append(f"Period: {period_str}")
            
            associated_works = entry.get("associated_works", [])
            if associated_works:
                lines.append("Associated Works:")
                for work in associated_works:
                    title = work.get("title_original", "")
                    role = work.get("role", "")
                    lines.append(f" - {title} ({role})")
            
            associated_places = entry.get("associated_places", [])
            if associated_places:
                lines.append("Associated Places:")
                for place in associated_places:
                    place_name = place.get("place_original", "")
                    assoc_type = place.get("association_type", "")
                    lines.append(f" - {place_name} ({assoc_type})")
            
            notes = entry.get("notes")
            if notes:
                lines.append(f"Notes: {notes}")
            
            lines.append("\n" + "=" * 40 + "\n")
        return lines

    def _convert_culinary_places_to_txt(self, entries: List[Any]) -> List[str]:
        """Converts culinary places entries to TXT format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        lines: List[str] = []
        for entry in entries:
            name = entry.get("name_original", "Unknown")
            lines.append(name)
            
            modern_name = entry.get("name_modern_english")
            if modern_name and modern_name != name:
                lines.append(f"Modern Name: {modern_name}")
            
            place_type = entry.get("place_type")
            country = entry.get("country_modern")
            if place_type:
                lines.append(f"Type: {place_type}")
            if country:
                lines.append(f"Country: {country}")
            
            period = entry.get("period", {})
            if period:
                period_str = f"{period.get('start_year', 'Unknown')} - {period.get('end_year', 'Unknown')}"
                if period.get("notation"):
                    period_str += f" ({period['notation']})"
                lines.append(f"Period: {period_str}")
            
            roles = entry.get("roles_in_culinary_ecosystem", [])
            if roles:
                lines.append(f"Roles: {', '.join(roles)}")
            
            products = entry.get("associated_products", [])
            if products:
                lines.append(f"Associated Products: {', '.join(products)}")
            
            establishments = entry.get("notable_establishments", [])
            if establishments:
                lines.append(f"Notable Establishments: {', '.join(establishments)}")
            
            notes = entry.get("notes")
            if notes:
                lines.append(f"Notes: {notes}")
            
            lines.append("\n" + "=" * 40 + "\n")
        return lines

    def _convert_culinary_works_to_txt(self, entries: List[Any]) -> List[str]:
        """Converts culinary works entries to TXT format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        lines: List[str] = []
        for entry in entries:
            title = entry.get("title_original", "Unknown")
            lines.append(title)
            
            modern_title = entry.get("title_modern_english")
            if modern_title and modern_title != title:
                lines.append(f"Modern Title: {modern_title}")
            
            short_title = entry.get("short_title")
            if short_title:
                lines.append(f"Short Title: {short_title}")
            
            genre = entry.get("genre")
            if genre:
                lines.append(f"Genre: {genre}")
            
            description = entry.get("description")
            if description:
                lines.append(f"Description: {description}")
            
            culinary_focus = entry.get("culinary_focus", [])
            if culinary_focus:
                lines.append(f"Culinary Focus: {', '.join(culinary_focus)}")
            
            languages = entry.get("languages", [])
            if languages:
                lines.append(f"Languages: {', '.join(languages)}")
            
            edition_years = entry.get("edition_years", [])
            if edition_years:
                years_str = ", ".join([str(y) for y in edition_years if y is not None])
                lines.append(f"Edition Years: {years_str}")
            
            contributors = entry.get("contributors", [])
            if contributors:
                lines.append("Contributors:")
                for contrib in contributors:
                    name = contrib.get("name_original", "")
                    role = contrib.get("role", "")
                    lines.append(f" - {name} ({role})")
            
            pub_places = entry.get("publication_places", [])
            if pub_places:
                place_names = [p.get("name_original", "") for p in pub_places]
                lines.append(f"Publication Places: {', '.join(place_names)}")
            
            notes = entry.get("notes")
            if notes:
                lines.append(f"Notes: {notes}")
            
            lines.append("\n" + "=" * 40 + "\n")
        return lines

    def _convert_historical_recipes_to_txt(self, entries: List[Any]) -> List[str]:
        """Converts historical recipes entries to TXT format."""
        # Filter out None entries
        if entries is None:
            entries = []
        entries = [entry for entry in entries if entry is not None]
        lines: List[str] = []
        for entry in entries:
            title = entry.get("title_original", "Unknown Recipe")
            lines.append(title)
            
            modern_title = entry.get("title_modern_english")
            if modern_title and modern_title != title:
                lines.append(f"Modern Title: {modern_title}")
            
            recipe_type = entry.get("recipe_type")
            if recipe_type:
                lines.append(f"Type: {recipe_type}")
            
            # Yield, prep time, cook time
            yields = entry.get("yield", [])
            if yields and len(yields) > 0:
                y = yields[0]
                val = y.get("value_modern_english")
                unit = y.get("unit_modern_english")
                if val and unit:
                    lines.append(f"Yield: {val} {unit}")
            
            prep_times = entry.get("preparation_time", [])
            if prep_times and len(prep_times) > 0:
                t = prep_times[0]
                val = t.get("value_modern_english")
                unit = t.get("unit_modern_english")
                if val and unit:
                    lines.append(f"Preparation Time: {val} {unit}")
            
            cook_times = entry.get("cooking_time", [])
            if cook_times and len(cook_times) > 0:
                t = cook_times[0]
                val = t.get("value_modern_english")
                unit = t.get("unit_modern_english")
                if val and unit:
                    lines.append(f"Cooking Time: {val} {unit}")
            
            # Ingredients
            ingredients = entry.get("ingredients", [])
            if ingredients:
                lines.append("Ingredients:")
                for ing in ingredients:
                    name = ing.get("name_modern_english") or ing.get("name_original") or ""
                    qty = ing.get("quantity_original") or ""
                    prep = ing.get("preparation_note_modern_english") or ""
                    ing_text = f" - {name}"
                    if qty:
                        ing_text += f" ({qty})"
                    if prep:
                        ing_text += f" - {prep}"
                    lines.append(ing_text)
            
            # Cooking methods
            methods = entry.get("cooking_methods", [])
            if methods:
                method_names = [m.get("method_modern_english") or m.get("method_original") or "" for m in methods]
                lines.append(f"Cooking Methods: {', '.join(method_names)}")
            
            # Original recipe text
            recipe_text = entry.get("recipe_text_original")
            if recipe_text:
                lines.append("Original Recipe Text:")
                lines.append(recipe_text)
            
            # Modern translation
            recipe_text_modern = entry.get("recipe_text_modern_english")
            if recipe_text_modern and recipe_text_modern != recipe_text:
                lines.append("Modern English Translation:")
                lines.append(recipe_text_modern)
            
            lines.append("\n" + "=" * 40 + "\n")
        return lines

    # --- Michelin Guide Converters ---

    def _convert_michelin_guides_to_docx(self, entries: List[Any], document: Document) -> None:
        """Convert Michelin Guide entries to DOCX format."""
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            
            # Header with establishment name and stars
            name = entry.get("establishment_name", "Unknown Establishment")
            awards = entry.get("awards", {}) or {}
            stars = awards.get("stars", 0) or 0
            star_display = "⭐" * stars if stars else ""
            
            document.add_heading(f"{name} {star_display}", level=1)
            
            # Location and Address
            location = entry.get("location", {}) or {}
            address = entry.get("address", {}) or {}
            city = location.get("city_or_town", "")
            neighbourhood = location.get("neighbourhood_or_area", "")
            street = address.get("street", "")
            house_number = address.get("house_number", "")
            postal_code = address.get("postal_code", "")
            
            location_parts = [p for p in [neighbourhood, city] if p]
            if location_parts:
                document.add_paragraph(f"Location: {', '.join(location_parts)}")
            
            address_parts = [p for p in [street, house_number, postal_code] if p]
            if address_parts:
                document.add_paragraph(f"Address: {' '.join(address_parts)}")
            
            # Awards section
            award_items = []
            if awards.get("bib_gourmand"):
                award_items.append("Bib Gourmand")
            if awards.get("green_star"):
                award_items.append("Green Star")
            if awards.get("michelin_plate"):
                award_items.append("Michelin Plate")
            if awards.get("new_in_guide"):
                award_items.append("New in Guide")
            if award_items:
                document.add_paragraph(f"Awards: {', '.join(award_items)}")
            
            # Cuisine
            cuisine = entry.get("cuisine", {}) or {}
            styles = cuisine.get("styles", [])
            if styles and isinstance(styles, list):
                document.add_paragraph(f"Cuisine: {', '.join(styles)}")
            chef = cuisine.get("chef")
            if chef:
                document.add_paragraph(f"Chef: {chef}")
            specialties = cuisine.get("specialties", [])
            if specialties and isinstance(specialties, list):
                document.add_paragraph(f"Specialties: {', '.join(specialties)}")
            
            # Pricing
            pricing = entry.get("pricing", {}) or {}
            currency = pricing.get("currency", "")
            menu_min = pricing.get("menu_price_min")
            menu_max = pricing.get("menu_price_max")
            if menu_min or menu_max:
                price_str = f"{currency} {menu_min or '?'} - {menu_max or '?'}"
                document.add_paragraph(f"Menu Price: {price_str}")
            
            alc_min = pricing.get("a_la_carte_price_min")
            alc_max = pricing.get("a_la_carte_price_max")
            if alc_min or alc_max:
                price_str = f"{currency} {alc_min or '?'} - {alc_max or '?'}"
                document.add_paragraph(f"À la carte: {price_str}")
            
            # Contact
            contact = entry.get("contact", {}) or {}
            tel = contact.get("telephone")
            if tel:
                document.add_paragraph(f"Telephone: {tel}")
            website = contact.get("website")
            if website:
                document.add_paragraph(f"Website: {website}")
            
            # Opening hours
            opening = entry.get("opening", {}) or {}
            lunch = opening.get("lunch_hours")
            dinner = opening.get("dinner_hours")
            if lunch:
                document.add_paragraph(f"Lunch: {lunch}")
            if dinner:
                document.add_paragraph(f"Dinner: {dinner}")
            days_closed = opening.get("days_closed", [])
            if days_closed and isinstance(days_closed, list):
                document.add_paragraph(f"Closed: {', '.join(days_closed)}")
            
            # Amenities
            amenities = entry.get("amenities", {}) or {}
            amenity_list = []
            if amenities.get("terrace"):
                amenity_list.append("Terrace")
            if amenities.get("garden_or_park"):
                amenity_list.append("Garden")
            if amenities.get("great_view"):
                amenity_list.append("Great View")
            if amenities.get("parking"):
                amenity_list.append("Parking")
            if amenities.get("wheelchair_access"):
                amenity_list.append("Wheelchair Access")
            if amenities.get("notable_wine_list"):
                amenity_list.append("Notable Wine List")
            if amenity_list:
                document.add_paragraph(f"Amenities: {', '.join(amenity_list)}")
            
            document.add_page_break()

    def _convert_michelin_guides_to_txt(self, entries: List[Any]) -> List[str]:
        """Convert Michelin Guide entries to TXT format."""
        lines: List[str] = []
        
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            
            # Header
            name = entry.get("establishment_name", "Unknown Establishment")
            awards = entry.get("awards", {}) or {}
            stars = awards.get("stars", 0) or 0
            star_display = "*" * stars if stars else "No stars"
            
            lines.append(f"{'=' * 60}")
            lines.append(f"{name}")
            lines.append(f"Stars: {star_display}")
            lines.append(f"{'=' * 60}")
            
            # Location and Address
            location = entry.get("location", {}) or {}
            address = entry.get("address", {}) or {}
            city = location.get("city_or_town", "")
            neighbourhood = location.get("neighbourhood_or_area", "")
            
            if city or neighbourhood:
                loc_str = ", ".join([p for p in [neighbourhood, city] if p])
                lines.append(f"Location: {loc_str}")
            
            street = address.get("street", "")
            house_number = address.get("house_number", "")
            postal_code = address.get("postal_code", "")
            if street or house_number or postal_code:
                addr_str = " ".join([p for p in [street, house_number, postal_code] if p])
                lines.append(f"Address: {addr_str}")
            
            # Awards
            award_items = []
            if awards.get("bib_gourmand"):
                award_items.append("Bib Gourmand")
            if awards.get("green_star"):
                award_items.append("Green Star")
            if awards.get("michelin_plate"):
                award_items.append("Michelin Plate")
            if awards.get("new_in_guide"):
                award_items.append("New")
            if award_items:
                lines.append(f"Awards: {', '.join(award_items)}")
            
            # Cuisine
            cuisine = entry.get("cuisine", {}) or {}
            styles = cuisine.get("styles", [])
            if styles and isinstance(styles, list):
                lines.append(f"Cuisine: {', '.join(styles)}")
            chef = cuisine.get("chef")
            if chef:
                lines.append(f"Chef: {chef}")
            specialties = cuisine.get("specialties", [])
            if specialties and isinstance(specialties, list):
                lines.append(f"Specialties: {', '.join(specialties)}")
            
            # Pricing
            pricing = entry.get("pricing", {}) or {}
            currency = pricing.get("currency", "")
            menu_min = pricing.get("menu_price_min")
            menu_max = pricing.get("menu_price_max")
            if menu_min or menu_max:
                lines.append(f"Menu: {currency} {menu_min or '?'} - {menu_max or '?'}")
            
            alc_min = pricing.get("a_la_carte_price_min")
            alc_max = pricing.get("a_la_carte_price_max")
            if alc_min or alc_max:
                lines.append(f"À la carte: {currency} {alc_min or '?'} - {alc_max or '?'}")
            
            # Contact
            contact = entry.get("contact", {}) or {}
            tel = contact.get("telephone")
            if tel:
                lines.append(f"Tel: {tel}")
            website = contact.get("website")
            if website:
                lines.append(f"Web: {website}")
            
            # Opening
            opening = entry.get("opening", {}) or {}
            lunch = opening.get("lunch_hours")
            dinner = opening.get("dinner_hours")
            if lunch:
                lines.append(f"Lunch: {lunch}")
            if dinner:
                lines.append(f"Dinner: {dinner}")
            days_closed = opening.get("days_closed", [])
            if days_closed and isinstance(days_closed, list):
                lines.append(f"Closed: {', '.join(days_closed)}")
            
            # Amenities
            amenities = entry.get("amenities", {}) or {}
            amenity_list = []
            if amenities.get("terrace"):
                amenity_list.append("Terrace")
            if amenities.get("garden_or_park"):
                amenity_list.append("Garden")
            if amenities.get("great_view"):
                amenity_list.append("View")
            if amenities.get("parking"):
                amenity_list.append("Parking")
            if amenities.get("wheelchair_access"):
                amenity_list.append("Wheelchair")
            if amenities.get("notable_wine_list"):
                amenity_list.append("Wine List")
            if amenity_list:
                lines.append(f"Amenities: {', '.join(amenity_list)}")
            
            lines.append("")
        
        return lines

    def _convert_cookbook_metadata_to_txt(self, entries: List[Any]) -> List[str]:
        """Convert cookbook metadata entries to the required plain text format."""
        lines: List[str] = []
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            lines.append(f"title: {self.safe_str(entry.get('title', 'unknown'))}")
            lines.append(f"author: {self.safe_str(entry.get('author', 'anonymous'))}")
            lines.append(f"year: {self.safe_str(entry.get('year', 'unknown'))}")
            lines.append(f"edition: {self.safe_str(entry.get('edition', 'unknown'))}")
            lines.append(f"content: {self.safe_str(entry.get('content', ''))}")
            lines.append(f"notes: {self.safe_str(entry.get('notes', ''))}")
            lines.append(f"library: {self.safe_str(entry.get('library', 'unknown'))}")
            lines.append(f"digitizer: {self.safe_str(entry.get('digitizer', 'unknown'))}")
            lines.append(f"misc: {self.safe_str(entry.get('misc', ''))}")
        return lines
