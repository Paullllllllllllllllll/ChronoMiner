"""Regression tests for the schema-v3.0 alignment of the conversion layer.

Covers the culinary Persons/Places/Works and unified Entities converters
against the shipped v3.0 schemas, the recipe production v3.0 analytic
columns, null tolerance in ``entries`` and optional list/object fields, and
the ``resolve_field`` explicit-null contract. Fixtures below are
schema-valid, nulls included.
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from modules.conversion.base import BaseConverter, resolve_field
from modules.conversion.csv_converter import CSVConverter
from modules.conversion.document_converter import DocumentConverter
from modules.conversion.json_utils import extract_entries_from_json


def _write_json(path: Path, payload: object) -> Path:
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    return path


def _entries_file(path: Path, entries: list) -> Path:
    return _write_json(
        path, {"contains_no_content_of_requested_type": False, "entries": entries}
    )


# ---------------------------------------------------------------------------
# Schema-valid fixtures (v3.0)
# ---------------------------------------------------------------------------

PERSON_V3 = {
    "names": {"original": "Bartolomeo Scappi", "modern_english": "Bartolomeo Scappi"},
    "short_notes": "Papal cook and author of the Opera.",
    "timeframe": {"start_year": 1540, "end_year": 1577, "notation": "fl. 1540-1577"},
    "lifespan": {"birth_year": 1500, "death_year": 1577},
    "geography": {
        "city_original": "Roma",
        "city_modern": "Rome",
        "country_original": "Stato Pontificio",
        "country_modern": "Italy",
    },
    "historical_importance": 7,
    "gender": "Male",
    "roles": ["Cook", "Author"],
    "associations": [
        {
            "entity_type": "Work",
            "entity_label_original": "Opera dell'arte del cucinare",
            "entity_label_modern": "The Opera of Bartolomeo Scappi",
            "relationship": "Created",
            "role": "Author",
        }
    ],
}

# Every optional field null — still schema-valid.
PERSON_V3_NULLS = {
    "names": {"original": "Anonymous Cook", "modern_english": None},
    "short_notes": None,
    "timeframe": None,
    "lifespan": None,
    "geography": None,
    "historical_importance": None,
    "gender": None,
    "roles": None,
    "associations": None,
}

PLACE_V3 = {
    "names": {"original": "Firenze", "modern_english": "Florence"},
    "short_notes": "Renaissance culinary center.",
    "timeframe": {"start_year": 1450, "end_year": 1600, "notation": "15th-16th c."},
    "events": [
        {"event_type": "Establishment", "year": 1400, "description": "Guild founded."}
    ],
    "geography": {
        "city_original": "Firenze",
        "city_modern": "Florence",
        "country_original": None,
        "country_modern": "Italy",
    },
    "historical_importance": 6,
    "place_type": "City",
    "roles_in_culinary_ecosystem": ["MarketCenter", "ProductionArea"],
    "associations": [
        {
            "entity_type": "Person",
            "entity_label_original": None,
            "entity_label_modern": "Catherine de' Medici",
            "relationship": "Birthplace",
            "role": None,
        }
    ],
}

PLACE_V3_NULLS = {
    "names": {"original": "Unnamed Market", "modern_english": None},
    "short_notes": None,
    "timeframe": None,
    "events": None,
    "geography": None,
    "historical_importance": None,
    "place_type": None,
    "roles_in_culinary_ecosystem": None,
    "associations": None,
}

WORK_V3 = {
    "titles": {
        "original": "Le Cuisinier François",
        "modern_english": "The French Cook",
        "short": "Cuisinier",
    },
    "short_notes": "Foundational text of French haute cuisine.",
    "timeframe": {"start_year": 1651, "end_year": 1700, "notation": "1651-1700"},
    "geography": {
        "city_original": "Paris",
        "city_modern": "Paris",
        "country_original": "France",
        "country_modern": "France",
    },
    "historical_importance": 7,
    "genre": "Cookbook",
    "culinary_focus": ["General", "Sauces and Condiments"],
    "languages": ["French"],
    "contributors": [
        {
            "name_original": "François Pierre de La Varenne",
            "name_modern_english": "La Varenne",
            "role": "Author",
        }
    ],
    "edition_years": [1651, 1652],
    "associations": [
        {
            "entity_type": "Place",
            "entity_label_original": "Paris",
            "entity_label_modern": "Paris",
            "relationship": "PublicationPlace",
            "role": None,
        }
    ],
}

WORK_V3_NULLS = {
    "titles": {"original": "Untitled", "modern_english": None, "short": None},
    "short_notes": None,
    "timeframe": None,
    "geography": None,
    "historical_importance": None,
    "genre": None,
    "culinary_focus": None,
    "languages": None,
    "contributors": None,
    "edition_years": None,
    "associations": None,
}

ENTITY_PERSON_V3 = {
    "entry_type": "Person",
    "person_entry": {
        "names": {"original": "Marie-Antoine Carême", "modern_english": "Carême"},
        "importance": 7,
        "summary": "Founder of French grande cuisine.",
        "timeframe": {
            "start_year": 1800,
            "end_year": 1833,
            "notation": "early 19th c.",
        },
        "geography": {"primary_location": "Paris", "additional_context": "France"},
        "topical_focus": ["Baking and Pastry"],
        "language_contexts": ["French"],
        "roles": ["ChefDeCuisine", "Author"],
    },
}

ENTITY_PLACE_V3 = {
    "entry_type": "Place",
    "place_entry": {
        "names": {"original": "Les Halles", "modern_english": "Les Halles"},
        "importance": 5,
        "summary": "Central Paris market.",
        "timeframe": None,
        "geography": {"primary_location": "Paris", "additional_context": None},
        "topical_focus": None,
        "language_contexts": None,
        "roles_in_culinary_ecosystem": ["MarketCenter"],
        "associated_products": ["vegetables"],
        "notable_establishments": ["Au Pied de Cochon"],
        "place_notes": "Demolished 1971.",
    },
}

ENTITY_WORK_V3 = {
    "entry_type": "Work",
    "work_entry": {
        "names": {"original": "Le Guide Culinaire", "modern_english": None},
        "importance": 6,
        "summary": None,
        "timeframe": None,
        "topical_focus": None,
        "language_contexts": None,
        "short_title": "Guide",
        "genre": "Cookbook",
    },
}

RECIPE_V3 = {
    "recipe_text_original": "Prenez du sucre...",
    "recipe_text_modern_english": "Take sugar...",
    "title_original": "Tarte au sucre",
    "title_modern_english": "Sugar tart",
    "recipe_type": "Pastry",
    "ingredients": [
        {
            "name_original": "sucre",
            "name_modern_english": "sugar",
            "quantity_original": "2 onces",
            "origin_explicitly_stated": "Indes",
            "ingredient_luxury_signal_rating_1_7": 5,
            "ingredient_trade_distance_rating_1_7": 6,
            "ingredient_novelty_rating_1_7": 4,
        },
        {
            "name_original": "beurre",
            "name_modern_english": "butter",
            "quantity_original": None,
            "origin_explicitly_stated": None,
            "ingredient_luxury_signal_rating_1_7": None,
            "ingredient_trade_distance_rating_1_7": 1,
            "ingredient_novelty_rating_1_7": 1,
        },
    ],
    "cooking_methods": [
        {
            "method_original": "cuire au four",
            "method_modern_english": "bake",
            "method_complexity_rating_1_7": 3,
        }
    ],
    "utensils_equipment": [
        {
            "utensil_original": "tourtière",
            "utensil_modern_english": "tart pan",
            "utensil_specialization_rating_1_7": 5,
            "utensil_modernity_rating_1_7": 3,
        },
        {
            "utensil_original": "four",
            "utensil_modern_english": "oven",
            "utensil_specialization_rating_1_7": 2,
            "utensil_modernity_rating_1_7": 2,
        },
    ],
    "timing_yield": {
        "yield_original": "six parts",
        "preparation_time_original": None,
        "cooking_time_original": "une heure",
    },
    "ingredient_categories": {"contains_refined_sugar": True, "contains_butter": True},
    "culinary_style": {
        "modernity_rating_1_7": 4,
        "innovation_markers_observed": ["Refined sugar use"],
        "archaism_markers_observed": [],
    },
    "intertextuality": {
        "explicit_source_attribution": None,
        "explicit_foreign_style_reference": None,
        "self_positioning_temporal": None,
        "tradition_claim_present": False,
        "authenticity_claim_present": False,
        "national_identity_claim_present": False,
        "anti_foreign_sentiment_present": False,
    },
    "geographic_signals": {
        "place_references": [
            {"place_name_original": "Indes", "reference_function": "Ingredient origin"}
        ]
    },
    "economic_signals": {
        "economic_framing_detected": ["Luxury/refinement framing"],
        "luxury_intensity_rating_1_7": 5,
        "occasion_type": ["Feast/banquet"],
    },
    "religious_signals": {
        "fasting_context_indicated": False,
        "meat_day_context_indicated": True,
        "confessional_hint": "No confessional signal",
        "moral_virtue_framing_present": False,
    },
}


# ---------------------------------------------------------------------------
# FIX 1 — culinary Persons / Places / Works read the v3.0 nested keys
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_culinary_persons_csv_reads_v3_keys(tmp_path: Path) -> None:
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryPersonsEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [PERSON_V3]), out
    )
    text = out.read_text(encoding="utf-8-sig")
    header = text.splitlines()[0]

    for column in (
        "name_original",
        "name_modern_english",
        "short_notes",
        "historical_importance",
        "timeframe_start_year",
        "birth_year",
        "death_year",
        "city_modern",
        "country_modern",
        "associations",
    ):
        assert column in header
    # Retired v2 columns must be gone.
    assert "canonical_name_original" not in header
    assert "period_start_year" not in header
    assert "links" not in header

    assert "Bartolomeo Scappi" in text
    assert "Rome" in text
    assert "1540" in text
    # Associations render the modern label.
    assert "The Opera of Bartolomeo Scappi" in text


@pytest.mark.unit
def test_culinary_persons_csv_survives_all_null_optionals(tmp_path: Path) -> None:
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryPersonsEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [PERSON_V3_NULLS]), out
    )
    text = out.read_text(encoding="utf-8-sig")
    assert "name_original" in text.splitlines()[0]
    assert "Anonymous Cook" in text


@pytest.mark.unit
def test_culinary_places_csv_reads_v3_keys(tmp_path: Path) -> None:
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryPlacesEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [PLACE_V3, PLACE_V3_NULLS]), out
    )
    text = out.read_text(encoding="utf-8-sig")
    header = text.splitlines()[0]

    for column in ("name_original", "events", "historical_importance", "associations"):
        assert column in header
    assert "associated_people" not in header

    assert "Firenze" in text
    assert "Florence" in text
    assert "MarketCenter" in text
    assert "Catherine de' Medici" in text
    assert "Guild founded." in text


@pytest.mark.unit
def test_culinary_works_csv_reads_v3_titles(tmp_path: Path) -> None:
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryWorksEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [WORK_V3, WORK_V3_NULLS]), out
    )
    text = out.read_text(encoding="utf-8-sig")
    header = text.splitlines()[0]

    for column in (
        "title_original",
        "title_modern_english",
        "title_short",
        "short_notes",
        "historical_importance",
        "associations",
    ):
        assert column in header
    assert "publication_places" not in header
    assert "description" not in header

    assert "Le Cuisinier François" in text
    assert "The French Cook" in text
    assert "La Varenne (Author)" in text
    assert "1651" in text


@pytest.mark.unit
def test_culinary_persons_txt_reads_v3_keys(tmp_path: Path) -> None:
    out = tmp_path / "out.txt"
    DocumentConverter("CulinaryPersonsEntries").convert_to_txt(
        _entries_file(tmp_path / "in.json", [PERSON_V3, PERSON_V3_NULLS]), out
    )
    text = out.read_text(encoding="utf-8")

    assert "Bartolomeo Scappi" in text
    assert "Timeframe: 1540 - 1577 (fl. 1540-1577)" in text
    assert "Lifespan: 1500 - 1577" in text
    assert "Geography: Rome, Italy" in text
    assert "Roles: Cook, Author" in text
    assert "Historical Importance: 7/7" in text
    assert "The Opera of Bartolomeo Scappi" in text
    # The all-null entry still renders its heading without crashing.
    assert "Anonymous Cook" in text
    assert "Unknown" not in text.splitlines()[0]


@pytest.mark.unit
def test_culinary_places_and_works_txt_read_v3_keys(tmp_path: Path) -> None:
    places_out = tmp_path / "places.txt"
    DocumentConverter("CulinaryPlacesEntries").convert_to_txt(
        _entries_file(tmp_path / "places.json", [PLACE_V3, PLACE_V3_NULLS]), places_out
    )
    places_text = places_out.read_text(encoding="utf-8")
    assert "Firenze" in places_text
    assert "Modern Name: Florence" in places_text
    assert "Type: City" in places_text
    assert "Establishment 1400: Guild founded." in places_text

    works_out = tmp_path / "works.txt"
    DocumentConverter("CulinaryWorksEntries").convert_to_txt(
        _entries_file(tmp_path / "works.json", [WORK_V3, WORK_V3_NULLS]), works_out
    )
    works_text = works_out.read_text(encoding="utf-8")
    assert "Le Cuisinier François" in works_text
    assert "Short Title: Cuisinier" in works_text
    assert "La Varenne (Author)" in works_text
    assert "Untitled" in works_text


@pytest.mark.unit
def test_culinary_docx_paths_render_v3_entries(tmp_path: Path) -> None:
    """DOCX rendering must not fall back to plain paragraphs for v3 entries."""
    from docx import Document

    for schema, entries in (
        ("CulinaryPersonsEntries", [PERSON_V3, PERSON_V3_NULLS]),
        ("CulinaryPlacesEntries", [PLACE_V3, PLACE_V3_NULLS]),
        ("CulinaryWorksEntries", [WORK_V3, WORK_V3_NULLS]),
    ):
        out = tmp_path / f"{schema}.docx"
        DocumentConverter(schema).convert_to_docx(
            _entries_file(tmp_path / f"{schema}.json", entries), out
        )
        assert out.exists()
        text = "\n".join(p.text for p in Document(str(out)).paragraphs)
        assert "{'names'" not in text  # no raw-dict fallback
        assert "Unknown" not in text

    docx_text = "\n".join(
        p.text
        for p in Document(str(tmp_path / "CulinaryPersonsEntries.docx")).paragraphs
    )
    assert "Geography: Rome, Italy" in docx_text


# ---------------------------------------------------------------------------
# FIX 2 — CulinaryEntitiesEntries matches the v3.0 profile shapes
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_culinary_entities_csv_matches_v3_profiles(tmp_path: Path) -> None:
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryEntitiesEntries").convert_to_csv(
        _entries_file(
            tmp_path / "in.json",
            [ENTITY_PERSON_V3, ENTITY_PLACE_V3, ENTITY_WORK_V3],
        ),
        out,
    )
    text = out.read_text(encoding="utf-8-sig")
    columns = text.splitlines()[0].split(",")

    for column in (
        "importance",
        "summary",
        "geography_primary_location",
        "geography_additional_context",
        "person_roles",
        "place_notes",
        "work_short_title",
        "work_genre",
    ):
        assert column in columns

    # Retired v2 columns must be gone.
    for column in (
        "entity_summary",
        "associations",
        "person_gender",
        "person_name_variants",
        "person_biographical_notes",
        "place_type",
        "place_country_modern",
        "work_description",
        "work_edition_years",
        "work_material_format",
        "notes",
    ):
        assert column not in columns

    assert "Carême" in text
    assert "Founder of French grande cuisine." in text
    assert "Au Pied de Cochon" in text
    assert "Guide" in text


@pytest.mark.unit
def test_culinary_entities_txt_matches_v3_profiles(tmp_path: Path) -> None:
    out = tmp_path / "out.txt"
    DocumentConverter("CulinaryEntitiesEntries").convert_to_txt(
        _entries_file(
            tmp_path / "in.json",
            [ENTITY_PERSON_V3, ENTITY_PLACE_V3, ENTITY_WORK_V3],
        ),
        out,
    )
    text = out.read_text(encoding="utf-8")

    assert "Importance: 7" in text
    assert "Summary: Founder of French grande cuisine." in text
    assert "Primary Location: Paris" in text
    assert "Roles: ChefDeCuisine, Author" in text
    assert "Place Notes: Demolished 1971." in text
    assert "Short Title: Guide" in text
    # Retired labels are gone.
    assert "Gender:" not in text
    assert "Name Variants:" not in text
    assert "Biographical Notes:" not in text
    assert "Material Notes:" not in text


# ---------------------------------------------------------------------------
# FIX 3 — a null ``entries`` value is schema-valid and must not raise
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_null_entries_top_level(tmp_path: Path) -> None:
    f = _write_json(
        tmp_path / "null.json",
        {"contains_no_content_of_requested_type": False, "entries": None},
    )
    assert extract_entries_from_json(f) == []


@pytest.mark.unit
def test_null_entries_in_one_record_does_not_kill_the_file(tmp_path: Path) -> None:
    payload = {
        "records": [
            {
                "custom_id": "chunk-1",
                "response": {
                    "contains_no_content_of_requested_type": False,
                    "entries": None,
                },
            },
            {"custom_id": "chunk-2", "response": {"entries": [{"id": 2}]}},
        ]
    }
    f = _write_json(tmp_path / "records.json", payload)
    assert extract_entries_from_json(f) == [{"id": 2}]


@pytest.mark.unit
def test_null_entries_in_serialized_response_text(tmp_path: Path) -> None:
    payload = {
        "records": [
            {
                "response": json.dumps(
                    {"contains_no_content_of_requested_type": False, "entries": None}
                )
            },
            {"response": json.dumps({"entries": [{"id": 7}]})},
        ]
    }
    f = _write_json(tmp_path / "records.json", payload)
    assert extract_entries_from_json(f) == [{"id": 7}]


@pytest.mark.unit
def test_null_entries_csv_still_writes_a_file(tmp_path: Path) -> None:
    f = _write_json(
        tmp_path / "null.json",
        {"contains_no_content_of_requested_type": False, "entries": None},
    )
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryPersonsEntries").convert_to_csv(f, out)
    assert out.exists()
    assert "name_original" in out.read_text(encoding="utf-8-sig").splitlines()[0]


# ---------------------------------------------------------------------------
# FIX 5 — recipe production v3.0 analytic columns
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_recipes_v3_emits_signal_columns(tmp_path: Path) -> None:
    out = tmp_path / "out.csv"
    CSVConverter("HistoricalRecipesEntriesProductionV3").convert_to_csv(
        _entries_file(tmp_path / "in.json", [RECIPE_V3]), out
    )
    text = out.read_text(encoding="utf-8-sig")
    header = text.splitlines()[0]

    for column in (
        "ingredient_origins_explicitly_stated",
        "utensil_specialization_ratings",
        "utensil_modernity_ratings",
        "place_references",
        "economic_framing_detected",
        "luxury_intensity_rating_1_7",
        "occasion_type",
        "fasting_context_indicated",
        "meat_day_context_indicated",
        "confessional_hint",
        "moral_virtue_framing_present",
    ):
        assert column in header

    assert "Indes (Ingredient origin)" in text
    assert "Luxury/refinement framing" in text
    assert "Feast/banquet" in text
    assert "No confessional signal" in text


@pytest.mark.unit
def test_recipes_v3_ratings_stay_index_parallel(tmp_path: Path) -> None:
    import csv

    out = tmp_path / "out.csv"
    CSVConverter("HistoricalRecipesEntriesProductionV3").convert_to_csv(
        _entries_file(tmp_path / "in.json", [RECIPE_V3]), out
    )
    with out.open(encoding="utf-8-sig", newline="") as fh:
        row = next(iter(csv.DictReader(fh)))

    assert row["ingredient_origins_explicitly_stated"] == "Indes; "
    assert row["utensils_equipment"] == "tart pan; oven"
    assert row["utensil_specialization_ratings"] == "5; 2"
    assert row["utensil_modernity_ratings"] == "3; 2"
    assert len(row["ingredients"].split("; ")) == 2


@pytest.mark.unit
def test_recipes_v3_documents_render_utensils(tmp_path: Path) -> None:
    from docx import Document

    json_file = _entries_file(tmp_path / "in.json", [RECIPE_V3])

    txt_out = tmp_path / "out.txt"
    DocumentConverter("HistoricalRecipesEntriesProductionV3").convert_to_txt(
        json_file, txt_out
    )
    txt = txt_out.read_text(encoding="utf-8")
    assert "Utensils and Equipment:" in txt
    assert "tart pan [Specialization: 5, Modernity: 3]" in txt

    docx_out = tmp_path / "out.docx"
    DocumentConverter("HistoricalRecipesEntriesProductionV3").convert_to_docx(
        json_file, docx_out
    )
    docx_text = "\n".join(p.text for p in Document(str(docx_out)).paragraphs)
    assert "tart pan [Specialization: 5, Modernity: 3]" in docx_text


# ---------------------------------------------------------------------------
# FIX 6 — a null list-of-dicts field must not flip the file to the fallback
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_null_list_field_keeps_the_dedicated_column_set(tmp_path: Path) -> None:
    entry = dict(PLACE_V3)
    entry["events"] = None
    entry["associations"] = None

    out = tmp_path / "out.csv"
    CSVConverter("CulinaryPlacesEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [entry]), out
    )
    header = out.read_text(encoding="utf-8-sig").splitlines()[0]

    # Dedicated-converter columns, not the json_normalize fallback's.
    assert "timeframe_notation" in header
    assert "names_original" not in header


# ---------------------------------------------------------------------------
# FIX 7 — a hostile entry degrades the DOCX/TXT paths instead of aborting
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_docx_and_txt_survive_non_dict_entries(tmp_path: Path) -> None:
    json_file = _entries_file(tmp_path / "in.json", ["not a dict", PERSON_V3])

    txt_out = tmp_path / "out.txt"
    DocumentConverter("CulinaryPersonsEntries").convert_to_txt(json_file, txt_out)
    assert txt_out.exists()
    assert "Bartolomeo Scappi" in txt_out.read_text(encoding="utf-8")

    docx_out = tmp_path / "out.docx"
    DocumentConverter("CulinaryPersonsEntries").convert_to_docx(json_file, docx_out)
    assert docx_out.exists()


# ---------------------------------------------------------------------------
# FIX 8 — resolve_field returns the default for explicit nulls
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_resolve_field_explicit_null_returns_default() -> None:
    assert resolve_field({"a": None}, "a", default="N/A") == "N/A"
    assert resolve_field({"a": {"b": None}}, "a.b", default="N/A") == "N/A"
    assert resolve_field({"a": None}, "a.b", default="N/A") == "N/A"
    # Falsy-but-present values are preserved.
    assert resolve_field({"a": 0}, "a", default="N/A") == 0
    assert resolve_field({"a": ""}, "a", default="N/A") == ""


@pytest.mark.unit
def test_addressbook_docx_does_not_print_none(tmp_path: Path) -> None:
    from docx import Document

    entry = {
        "last_name": "Meier",
        "first_name": "Hans",
        "occupation": "Baker",
        "address": {"street": None, "street_number": None},
        "honorific": None,
        "additional_notes": None,
    }
    json_file = _entries_file(tmp_path / "in.json", [entry])

    docx_out = tmp_path / "out.docx"
    DocumentConverter("HistoricalAddressBookEntries").convert_to_docx(
        json_file, docx_out
    )
    text = "\n".join(p.text for p in Document(str(docx_out)).paragraphs)
    assert "None" not in text
    assert "*0*" not in text
    assert "Address: Unknown" in text


@pytest.mark.unit
def test_cookbook_txt_and_docx_agree_on_nulls(tmp_path: Path) -> None:
    from docx import Document

    entry = {
        "title": None,
        "author": None,
        "year": None,
        "edition": None,
        "content": None,
        "notes": None,
        "library": None,
        "digitizer": None,
        "misc": None,
    }
    json_file = _entries_file(
        tmp_path / "in.json", [entry, dict(entry, title="Second")]
    )

    txt_out = tmp_path / "out.txt"
    DocumentConverter("CookbookMetadataEntries").convert_to_txt(json_file, txt_out)
    txt = txt_out.read_text(encoding="utf-8")
    assert "None" not in txt
    assert "author: anonymous" in txt
    # Entries are separated, as in every other TXT converter.
    assert "=" * 40 in txt

    docx_out = tmp_path / "out.docx"
    DocumentConverter("CookbookMetadataEntries").convert_to_docx(json_file, docx_out)
    docx_text = "\n".join(p.text for p in Document(str(docx_out)).paragraphs)
    assert "None" not in docx_text
    assert "Author: anonymous" in docx_text


@pytest.mark.unit
def test_recipe_title_null_falls_back(tmp_path: Path) -> None:
    entry = dict(RECIPE_V3, title_original=None, title_modern_english=None)
    out = tmp_path / "out.txt"
    DocumentConverter("HistoricalRecipesEntriesProductionV3").convert_to_txt(
        _entries_file(tmp_path / "in.json", [entry]), out
    )
    assert out.read_text(encoding="utf-8").splitlines()[0] == "Unknown Recipe"


# ---------------------------------------------------------------------------
# FIX 9 — the top-level-list fallback must not resurrect chunk records
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_top_level_list_of_empty_records_yields_no_entries(tmp_path: Path) -> None:
    payload = [
        {"custom_id": "chunk-1", "response": {"entries": []}},
        {
            "custom_id": "chunk-2",
            "response": {
                "contains_no_content_of_requested_type": True,
                "entries": None,
            },
        },
    ]
    f = _write_json(tmp_path / "chunks.json", payload)
    assert extract_entries_from_json(f) == []


@pytest.mark.unit
def test_top_level_list_of_bare_entries_still_falls_back(tmp_path: Path) -> None:
    payload = [{"full_title": "A Book"}, {"full_title": "Another Book"}]
    f = _write_json(tmp_path / "entries.json", payload)
    assert extract_entries_from_json(f) == payload


# ---------------------------------------------------------------------------
# FIX 10 — the retired MichelinGuides converters are gone, Light is intact
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_retired_michelin_guides_converters_removed() -> None:
    for cls in (CSVConverter, DocumentConverter):
        assert not [n for n in dir(cls) if "michelin_guides_to" in n]
    assert not hasattr(DocumentConverter, "_MICHELIN_AWARD_KEYS")
    assert not hasattr(DocumentConverter, "_MICHELIN_AMENITY_KEYS")


@pytest.mark.unit
def test_michelin_light_tolerates_string_star_counts(tmp_path: Path) -> None:
    entry = {
        "establishment_name": "Chez Stringy",
        "awards": {"stars": "2"},
        "location": {"city_or_town": "Lyon"},
    }
    json_file = _entries_file(tmp_path / "in.json", [entry])

    txt_out = tmp_path / "out.txt"
    DocumentConverter("MichelinGuidesLight").convert_to_txt(json_file, txt_out)
    assert "Stars: **" in txt_out.read_text(encoding="utf-8")

    docx_out = tmp_path / "out.docx"
    DocumentConverter("MichelinGuidesLight").convert_to_docx(json_file, docx_out)
    assert docx_out.exists()


# ---------------------------------------------------------------------------
# FIX 1 (base helpers) — timeframe default key and association labels
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_base_period_helpers_default_to_timeframe() -> None:
    entry = {"timeframe": {"start_year": 1600, "end_year": 1650, "notation": "17th c."}}
    assert BaseConverter._extract_period(entry) == (1600, 1650, "17th c.")
    assert BaseConverter._format_period(entry) == "1600 - 1650 (17th c.)"
    assert BaseConverter._extract_period({"timeframe": None}) == (None, None, None)
    assert BaseConverter._format_period({"timeframe": None}) == ""


@pytest.mark.unit
def test_format_links_prefers_modern_entity_label() -> None:
    links = [
        {
            "entity_type": "Work",
            "entity_label_original": "Opera",
            "entity_label_modern": "The Opera",
            "relationship": "Created",
        },
        {
            "entity_type": "Place",
            "entity_label_original": "Roma",
            "entity_label_modern": None,
            "relationship": "LocatedIn",
        },
    ]
    result = BaseConverter._format_links(links)
    assert "Work: The Opera - Created" in result
    assert "Place: Roma - LocatedIn" in result
    assert BaseConverter._format_links(None) == ""


# ---------------------------------------------------------------------------
# FIX 11 — CSV rendering defects: nullable ints, BOM, ragged null cells
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_nullable_int_columns_do_not_render_as_floats(tmp_path: Path) -> None:
    """A null in an integer column must not turn 1651 into '1651.0'."""
    import csv

    out = tmp_path / "out.csv"
    CSVConverter("CulinaryWorksEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [WORK_V3, WORK_V3_NULLS]), out
    )
    text = out.read_text(encoding="utf-8-sig")
    assert "1651.0" not in text
    assert "7.0" not in text

    with out.open(encoding="utf-8-sig", newline="") as fh:
        rows = list(csv.DictReader(fh))
    assert rows[0]["timeframe_start_year"] == "1651"
    assert rows[0]["historical_importance"] == "7"
    assert rows[1]["timeframe_start_year"] == ""
    # Strings are untouched by the dtype conversion.
    assert rows[0]["title_modern_english"] == "The French Cook"
    assert rows[1]["title_modern_english"] == ""


@pytest.mark.unit
def test_csv_is_written_with_a_utf8_bom(tmp_path: Path) -> None:
    """Excel needs the BOM to read the UTF-8 output correctly."""
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryPersonsEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [PERSON_V3]), out
    )
    assert out.read_bytes().startswith(b"\xef\xbb\xbf")
    # The BOM is the only difference: content still decodes cleanly.
    assert "Bartolomeo Scappi" in out.read_text(encoding="utf-8-sig")


@pytest.mark.unit
def test_format_links_drops_missing_parts() -> None:
    """Ragged association objects must not produce ': X - ' fragments."""
    assert (
        BaseConverter._format_links(
            [{"entity_type": None, "entity_label_modern": "Escoffier"}]
        )
        == "Escoffier"
    )
    assert (
        BaseConverter._format_links(
            [{"entity_type": "Person", "entity_label_modern": "Escoffier"}]
        )
        == "Person: Escoffier"
    )
    assert (
        BaseConverter._format_links(
            [{"entity_label_original": "Escoffier", "relationship": "StudentOf"}]
        )
        == "Escoffier - StudentOf"
    )
    assert BaseConverter._format_links([{"relationship": "StudentOf"}]) == "StudentOf"
    # An all-null association contributes no cell fragment at all.
    assert (
        BaseConverter._format_links(
            [
                {
                    "entity_type": None,
                    "entity_label_modern": None,
                    "entity_label_original": None,
                    "relationship": None,
                }
            ]
        )
        == ""
    )


@pytest.mark.unit
def test_event_cells_omit_empty_parens_and_colons(tmp_path: Path) -> None:
    import csv

    entry = dict(PLACE_V3)
    entry["events"] = [
        {"event_type": "Establishment", "year": 1400, "description": "Guild founded."},
        {"event_type": "Fire", "year": None, "description": None},
        {"event_type": None, "year": None, "description": None},
    ]
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryPlacesEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [entry]), out
    )
    with out.open(encoding="utf-8-sig", newline="") as fh:
        row = next(iter(csv.DictReader(fh)))

    assert row["events"] == "Establishment (1400): Guild founded.; Fire"
    assert "()" not in row["events"]


@pytest.mark.unit
def test_contributor_cells_omit_empty_parens(tmp_path: Path) -> None:
    import csv

    entry = dict(WORK_V3)
    entry["contributors"] = [
        {"name_original": "La Varenne", "name_modern_english": None, "role": "Author"},
        {"name_original": None, "name_modern_english": None, "role": "Editor"},
        {"name_original": "Pierre David", "name_modern_english": None, "role": None},
        {"name_original": None, "name_modern_english": None, "role": None},
    ]
    out = tmp_path / "out.csv"
    CSVConverter("CulinaryWorksEntries").convert_to_csv(
        _entries_file(tmp_path / "in.json", [entry]), out
    )
    with out.open(encoding="utf-8-sig", newline="") as fh:
        row = next(iter(csv.DictReader(fh)))

    assert row["contributors"] == "La Varenne (Author); Editor; Pierre David"
    assert " ()" not in row["contributors"]


# ---------------------------------------------------------------------------
# FIX 12 — schema-valid nulls must never surface as the literal "None"
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_addressbook_header_null_occupation(tmp_path: Path) -> None:
    entry = {
        "last_name": "Meier",
        "first_name": "Hans",
        "occupation": None,
        "section": None,
    }
    out = tmp_path / "out.txt"
    DocumentConverter("HistoricalAddressBookEntries").convert_to_txt(
        _entries_file(tmp_path / "in.json", [entry]), out
    )
    text = out.read_text(encoding="utf-8")
    assert "None" not in text
    assert "Meier, Hans - Unknown" in text


@pytest.mark.unit
def test_brazilian_header_null_profession(tmp_path: Path) -> None:
    entry = {"surname": "Silva", "first_name": "Joao", "profession": None}
    out = tmp_path / "out.txt"
    DocumentConverter("BrazilianMilitaryRecords").convert_to_txt(
        _entries_file(tmp_path / "in.json", [entry]), out
    )
    assert "None" not in out.read_text(encoding="utf-8")


@pytest.mark.unit
def test_format_officials_drops_null_signature() -> None:
    assert (
        BaseConverter._format_officials(
            {"officials": [{"position": "Captain", "signature": None}]}
        )
        == "Captain"
    )
    assert (
        BaseConverter._format_officials(
            {"officials": [{"position": None, "signature": "J. Silva"}]}
        )
        == "J. Silva"
    )
    assert (
        BaseConverter._format_officials(
            {"officials": [{"position": "Captain", "signature": "J. Silva"}]}
        )
        == "Captain: J. Silva"
    )
    assert (
        BaseConverter._format_officials(
            {"officials": [{"position": None, "signature": None}]}
        )
        == ""
    )


@pytest.mark.unit
def test_bibliographic_txt_and_docx_agree_on_null_edition_fields(
    tmp_path: Path,
) -> None:
    from docx import Document

    entry = {
        "full_title": None,
        "short_title": None,
        "main_author": None,
        "edition_info": [
            {
                "year": None,
                "edition_number": None,
                "publication_locations": None,
                "contributors": [
                    {"name": None, "role": "Editor"},
                    {"name": "Pierre David", "role": None},
                    {"name": None, "role": None},
                ],
            }
        ],
    }
    json_file = _entries_file(tmp_path / "in.json", [entry])

    txt_out = tmp_path / "out.txt"
    DocumentConverter("BibliographicEntries").convert_to_txt(json_file, txt_out)
    txt = txt_out.read_text(encoding="utf-8")

    docx_out = tmp_path / "out.docx"
    DocumentConverter("BibliographicEntries").convert_to_docx(json_file, docx_out)
    docx_text = "\n".join(p.text for p in Document(str(docx_out)).paragraphs)

    for text in (txt, docx_text):
        assert "None" not in text
        assert "Unknown Title" in text
        assert "Year: Unknown" in text
        assert "Edition: Unknown" in text
        assert "Editor, Pierre David" in text
        assert " ()" not in text


# ---------------------------------------------------------------------------
# FIX 13 — hostile non-dict elements degrade locally, not file-wide
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_brazilian_documents_survive_non_dict_entries(tmp_path: Path) -> None:
    from docx import Document

    json_file = _entries_file(
        tmp_path / "in.json", ["not a dict", {"surname": "Silva", "first_name": "Joao"}]
    )

    txt_out = tmp_path / "out.txt"
    DocumentConverter("BrazilianMilitaryRecords").convert_to_txt(json_file, txt_out)
    assert "Silva" in txt_out.read_text(encoding="utf-8")

    docx_out = tmp_path / "out.docx"
    DocumentConverter("BrazilianMilitaryRecords").convert_to_docx(json_file, docx_out)
    docx_text = "\n".join(p.text for p in Document(str(docx_out)).paragraphs)
    assert "Record Header" in docx_text


@pytest.mark.unit
def test_recipe_documents_survive_non_dict_list_elements(tmp_path: Path) -> None:
    from docx import Document

    entry = dict(RECIPE_V3)
    entry["ingredients"] = ["hostile", *RECIPE_V3["ingredients"]]
    entry["cooking_methods"] = ["hostile", *RECIPE_V3["cooking_methods"]]
    json_file = _entries_file(tmp_path / "in.json", [entry])

    txt_out = tmp_path / "out.txt"
    DocumentConverter("HistoricalRecipesEntriesProductionV3").convert_to_txt(
        json_file, txt_out
    )
    txt = txt_out.read_text(encoding="utf-8")
    assert "Tarte au sucre" in txt
    assert "sugar" in txt

    docx_out = tmp_path / "out.docx"
    DocumentConverter("HistoricalRecipesEntriesProductionV3").convert_to_docx(
        json_file, docx_out
    )
    docx_text = "\n".join(p.text for p in Document(str(docx_out)).paragraphs)
    assert "sugar" in docx_text


# ---------------------------------------------------------------------------
# FIX 14 — list-form Chat-Completions content must not abort the conversion
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_list_form_chat_completions_content_is_extracted(tmp_path: Path) -> None:
    payload = {
        "records": [
            {
                "custom_id": "chunk-1",
                "response": {
                    "choices": [
                        {
                            "message": {
                                "content": [
                                    {"type": "text", "text": '{"entries": '},
                                    {"type": "text", "text": '[{"id": 5}]}'},
                                    "hostile non-dict block",
                                ]
                            }
                        }
                    ]
                },
            }
        ]
    }
    f = _write_json(tmp_path / "records.json", payload)
    assert extract_entries_from_json(f) == [{"id": 5}]


@pytest.mark.unit
def test_non_string_chat_completions_content_yields_no_entries(tmp_path: Path) -> None:
    payload = {
        "records": [{"response": {"choices": [{"message": {"content": 42}}]}}],
    }
    f = _write_json(tmp_path / "records.json", payload)
    assert extract_entries_from_json(f) == []


# ---------------------------------------------------------------------------
# FIX 15 — the dead legacy formatting helpers are gone
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_legacy_base_converter_helpers_removed() -> None:
    for name in (
        "format_associations",
        "format_name_variants",
        "_extract_first_measurement",
    ):
        assert not hasattr(BaseConverter, name)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
